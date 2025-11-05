from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_ALIGN_VERTICAL
import docx

def beautify_table(table, header_rows=1):
        # 设置所有单元格左右、上下居中，黑色边框
        for row in table.rows:
            # 设置行高
            row.height = Cm(1.2)  # 设置行高为1.2厘米
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置数据行字体样式
                    if not p.runs:
                        # 如果没有runs，为现有文本添加run
                        p.add_run(cell.text)
                    for run in p.runs:
                        run.font.size = docx.shared.Pt(13)
                        run.font.name = 'Source Han Sans CN'
                        run.font.color.rgb = RGBColor(0, 0, 0)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 黑色边框
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="8" w:color="000000"/><w:left w:val="single" w:sz="8" w:color="000000"/><w:bottom w:val="single" w:sz="8" w:color="000000"/><w:right w:val="single" w:sz="8" w:color="000000"/></w:tcBorders>'.format(nsdecls('w'))))
        # 标题行美化
        for col in range(len(table.columns)):
            cell = table.cell(0, col)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
            run.font.bold = True
            run.font.size = docx.shared.Pt(13)
            run.font.name = 'Source Han Sans CN'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9EAFB"/>'.format(nsdecls('w'))))
            run.font.color.rgb = RGBColor(0, 0, 0)

def add_chart_title(doc, title_text):
    """为图表添加美化的标题"""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title_text)
    title_run.font.size = docx.shared.Pt(16)
    title_run.font.bold = True
    title_run.font.name = 'Source Han Sans CN'
    title_run.font.color.rgb = RGBColor(30, 64, 175)  # 深蓝色
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = docx.shared.Pt(12)
    title_para.paragraph_format.space_after = docx.shared.Pt(8)
    title_para.paragraph_format.left_indent = docx.shared.Pt(0)
    
    # 添加下划线装饰
    try:
        # 设置段落边框
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1F40AF"/></w:pBdr>' % nsdecls('w'))
        title_para._p.get_or_add_pPr().append(shading_elm)
    except:
        # 如果设置边框失败，使用备用方案
        title_para.paragraph_format.border_bottom = docx.shared.BorderStyle.SINGLE
    
    return title_para

def add_section_title(doc, title_text):
    """为1级标题添加美化的标题（比图表标题更大）"""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title_text)
    title_run.font.size = docx.shared.Pt(20)  # 比图表标题的16pt更大
    title_run.font.bold = True
    title_run.font.name = 'Source Han Sans CN'
    title_run.font.color.rgb = RGBColor(30, 64, 175)  # 深蓝色
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = docx.shared.Pt(16)
    title_para.paragraph_format.space_after = docx.shared.Pt(10)
    title_para.paragraph_format.left_indent = docx.shared.Pt(0)
    
    # 添加下划线装饰，比图表标题更粗
    try:
        # 设置段落边框
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="18" w:space="1" w:color="1F40AF"/></w:pBdr>' % nsdecls('w'))
        title_para._p.get_or_add_pPr().append(shading_elm)
    except:
        # 如果设置边框失败，使用备用方案
        title_para.paragraph_format.border_bottom = docx.shared.BorderStyle.SINGLE
    
    return title_para

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import io
import matplotlib.pyplot as plt
import matplotlib
import matplotlib.font_manager as fm
matplotlib.use('Agg')  # 使用非交互式后端

# 配置思源黑体字体支持
def setup_chinese_font():
    """设置中文字体支持，优先使用思源黑体"""
    # 固定使用思源黑体（不做优先级、不做验证、不做初始化判断）
    font_path = '/Users/yc/Documents/09_SourceHanSansSC/OTF/SimplifiedChinese/SourceHanSansSC-Regular.otf'
    fm.fontManager.addfont(font_path)
    prop = fm.FontProperties(fname=font_path)
    plt.rcParams['font.family'] = prop.get_name()
    plt.rcParams['axes.unicode_minus'] = False

# 初始化字体配置
setup_chinese_font()

def filter_device_count_data(df):
    """
    过滤用于设备数量统计的数据，排除资产状态列中包含'费'或'赔偿'的记录
    """
    if df.empty:
        return df
    
    # 如果没有资产状态列，直接返回原数据
    if '资产状态' not in df.columns:
        return df
    
    # 过滤条件：排除资产状态列中包含'费'或'赔偿'的行
    filter_condition = ~(
        df['资产状态'].astype(str).str.contains('费', na=False) |
        df['资产状态'].astype(str).str.contains('赔偿', na=False)
    )
    
    return df[filter_condition]

def ensure_chinese_font():
    """确保matplotlib中文字体正确设置（在每次生成图表前调用）"""
    setup_chinese_font()

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import tempfile
import base64

# 全局样式定义（整合PDF视觉元素）
st.markdown("""
<style>
    /* 全局表格居中强制规则 */
    div[data-testid="stDataFrame"] table td,
    div[data-testid="stDataFrame"] table th,
    .stDataFrame table td,
    .stDataFrame table th {
        text-align: center !important;
        vertical-align: middle !important;
    }
    div[data-testid="stDataFrame"] {
        text-align: center !important;
    }
    
    /* 更强制的表格居中规则 */
    [data-testid="stDataFrame"] * {
        text-align: center !important;
    }
    
    /* 针对具体的表格单元格内容 */
    [data-testid="stDataFrame"] tbody tr td div,
    [data-testid="stDataFrame"] thead tr th div {
        text-align: center !important;
        justify-content: center !important;
        align-items: center !important;
        display: flex !important;
        width: 100% !important;
    }
    
    /* 针对数据内容的强制居中 */
    [data-testid="stDataFrame"] td > div,
    [data-testid="stDataFrame"] th > div {
        text-align: center !important;
        justify-content: center !important;
        display: flex !important;
        width: 100% !important;
    }
    
    /* 主容器 */
    .main-container {
        max-width: 1200px;
        margin: 0 auto;
        padding: 20px 30px;
    }
    /* 标题样式（对应PDF中的💻 IT设备租赁分析系统） */
    .page-title {
        text-align: center;
        color: #2c3e50;
        font-size: 28px;
        font-weight: 700;
        margin: 10px 0 30px;
        padding-bottom: 15px;
        border-bottom: 3px solid #4a8cff;
    }
    .section-title {
        font-size: 20px;
        font-weight: 600;
        color: #2c3e50;
        margin: 25px 0 15px;
        padding-left: 10px;
        border-left: 4px solid #4a8cff;
    }
    /* 卡片样式（对应PDF中的数据区块） */
    .data-card {
        background: white;
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        padding: 20px;
        margin-bottom: 20px;
    }
    .metric-card {
        background: #f8f9fa;
        border-radius: 10px;
        padding: 18px;
        text-align: center;
        transition: transform 0.3s;
    }
    .metric-card:hover {
        transform: translateY(-3px);
    }
    .metric-value {
        font-size: 22px;
        font-weight: bold;
        color: #2c3e50;
    }
    .metric-label {
        font-size: 14px;
        color: #6c757d;
        margin-top: 5px;
    }
    /* 表格样式（对应PDF中的表格） */
    .dataframe-container {
        overflow-x: auto;
        border-radius: 8px;
        border: 1px solid #e9ecef;
    }
    /* Streamlit 表格样式优化 */
    .stDataFrame {
        text-align: center;
    }
    .stDataFrame table {
        margin: 0 auto;
        border-collapse: collapse;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    .stDataFrame th {
        background-color: #d9eafb !important;
        color: #000000 !important;
        font-weight: bold !important;
        text-align: center !important;
        vertical-align: middle !important;
        padding: 10px 8px !important;
        border: 1px solid #b3d9f5 !important;
        font-size: 14px !important;
        line-height: 1.4 !important;
        height: 40px !important;
    }
    .stDataFrame td {
        text-align: center !important;
        vertical-align: middle !important;
        padding: 8px !important;
        border: 1px solid #e9ecef !important;
        font-size: 13px !important;
        line-height: 1.4 !important;
        height: 35px !important;
    }
    
    /* 超强制居中规则 */
    .stDataFrame td *,
    .stDataFrame th * {
        text-align: center !important;
        justify-content: center !important;
        align-items: center !important;
    }
    
    /* 针对Streamlit特殊结构的居中 */
    .stDataFrame [data-testid="stDataFrame"] tbody tr td,
    .stDataFrame [data-testid="stDataFrame"] thead tr th {
        text-align: center !important;
    }
    
    /* 针对所有可能的div嵌套结构 */
    .stDataFrame div,
    .stDataFrame span,
    .stDataFrame p {
        text-align: center !important;
    }
    .stDataFrame tr:nth-child(even) {
        background-color: #f8f9fa !important;
    }
    .stDataFrame tr:hover {
        background-color: #e3f2fd !important;
    }
    /* 强制所有表格内容居中 */
    .stDataFrame [data-testid="stDataFrame"] {
        text-align: center !important;
    }
    .stDataFrame [data-testid="stDataFrame"] table {
        margin: 0 auto !important;
    }
    .stDataFrame [data-testid="stDataFrame"] td,
    .stDataFrame [data-testid="stDataFrame"] th {
        text-align: center !important;
        vertical-align: middle !important;
    }
    /* 增强表格内容居中效果 */
    .stDataFrame tbody tr td div {
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        height: 100% !important;
    }
    .stDataFrame thead tr th div {
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        height: 100% !important;
    }
    /* 进一步确保所有表格内容居中 */
    .stDataFrame * {
        text-align: center !important;
    }
    .stDataFrame tbody tr td {
        text-align: center !important;
        vertical-align: middle !important;
    }
    .stDataFrame thead tr th {
        text-align: center !important;
        vertical-align: middle !important;
    }
    /* 针对数据内容的特殊处理 */
    .stDataFrame tbody tr td > div > div {
        text-align: center !important;
        width: 100% !important;
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
    }
    
    /* 最强制的居中规则 - 覆盖所有可能的嵌套结构 */
    .stDataFrame table tbody tr td,
    .stDataFrame table thead tr th,
    .stDataFrame table tbody tr td > *,
    .stDataFrame table thead tr th > * {
        text-align: center !important;
        justify-content: center !important;
        align-items: center !important;
    }
    
    /* 针对具体的数据单元格 */
    .stDataFrame table tbody tr td div[data-testid],
    .stDataFrame table thead tr th div[data-testid] {
        text-align: center !important;
        justify-content: center !important;
        display: flex !important;
        width: 100% !important;
    }
    
    /* 强制所有文本内容居中 */
    .stDataFrame table * {
        text-align: center !important;
    }
    
    /* 针对Streamlit表格内部所有元素的超强制居中 */
    .stDataFrame [data-testid="stDataFrame"] tbody tr td > div,
    .stDataFrame [data-testid="stDataFrame"] thead tr th > div {
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        text-align: center !important;
        width: 100% !important;
        height: 100% !important;
    }
    
    /* 确保所有嵌套的div都居中 */
    .stDataFrame [data-testid="stDataFrame"] tbody tr td > div > div,
    .stDataFrame [data-testid="stDataFrame"] thead tr th > div > div {
        text-align: center !important;
        width: 100% !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    
    /* 针对表格内所有文本节点的强制居中 */
    .stDataFrame [data-testid="stDataFrame"] tbody tr td span,
    .stDataFrame [data-testid="stDataFrame"] thead tr th span,
    .stDataFrame [data-testid="stDataFrame"] tbody tr td p,
    .stDataFrame [data-testid="stDataFrame"] thead tr th p {
        text-align: center !important;
        width: 100% !important;
        display: block !important;
        margin: 0 auto !important;
    }
    
    /* 覆盖所有可能的表格内容结构 */
    .stDataFrame table tbody tr td > *,
    .stDataFrame table thead tr th > * {
        text-align: center !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        width: 100% !important;
        height: 100% !important;
    }
    
    /* 最终兜底规则 - 确保所有表格内容都居中 */
    .stDataFrame table tbody tr,
    .stDataFrame table thead tr {
        text-align: center !important;
    }
    
    .stDataFrame table tbody tr td,
    .stDataFrame table thead tr th {
        text-align: center !important;
        vertical-align: middle !important;
    }
    
    /* 针对数据单元格的特殊处理 */
    .stDataFrame table tbody tr td[data-testid],
    .stDataFrame table thead tr th[data-testid] {
        text-align: center !important;
        vertical-align: middle !important;
    }
    
    /* 确保表格的flex布局居中 */
    .stDataFrame [data-testid="stDataFrame"] {
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
        text-align: center !important;
    }
    
    .stDataFrame [data-testid="stDataFrame"] table {
        text-align: center !important;
        margin: 0 auto !important;
    }
    /* 图表样式（对应PDF中的柱状图/饼图） */
    .chart-container {
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }
    
    /* 追加：针对Streamlit最新版本的表格结构优化 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td > div[data-testid] {
        text-align: center !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        width: 100% !important;
        height: 100% !important;
    }
    
    /* 针对表格内部可能的文本包装元素 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td div[data-testid] > div,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th div[data-testid] > div {
        text-align: center !important;
        width: 100% !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    
    /* 确保表格行和列的flex布局居中 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr,
    .stDataFrame [data-testid="stDataFrame"] table thead tr {
        display: table-row !important;
        text-align: center !important;
    }
    
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th {
        display: table-cell !important;
        text-align: center !important;
        vertical-align: middle !important;
    }
    
    /* 针对可能的内联元素进行居中 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td > div > span,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th > div > span {
        text-align: center !important;
        display: inline-block !important;
        width: 100% !important;
    }
    
    /* 最终强制规则：覆盖所有可能的表格内容 */
    .stDataFrame [data-testid="stDataFrame"] table * {
        text-align: center !important;
    }
    
    /* 确保表格容器本身也居中 */
    .stDataFrame {
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
        width: 100% !important;
    }
    
    /* 终极居中规则 - 使用重要性更高的选择器 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th {
        text-align: center !important;
        vertical-align: middle !important;
        padding: 8px !important;
        font-size: 13px !important;
        line-height: 1.4 !important;
        border: 1px solid #e9ecef !important;
    }
    
    /* 确保所有可能的嵌套元素都居中 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td > div,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th > div,
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td > div > div,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th > div > div,
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td > div > div > div,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th > div > div > div {
        text-align: center !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        width: 100% !important;
        height: 100% !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    
    /* 文本元素的居中 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td span,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th span,
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td p,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th p,
    .stDataFrame [data-testid="stDataFrame"] table tbody tr td strong,
    .stDataFrame [data-testid="stDataFrame"] table thead tr th strong {
        text-align: center !important;
        width: 100% !important;
        display: block !important;
        margin: 0 auto !important;
        padding: 0 !important;
    }
    
    /* 超级强制规则 - 覆盖所有Streamlit内部样式 */
    .stDataFrame [data-testid="stDataFrame"] table * {
        text-align: center !important;
        vertical-align: middle !important;
    }
    
    /* 表格行的居中 */
    .stDataFrame [data-testid="stDataFrame"] table tbody tr,
    .stDataFrame [data-testid="stDataFrame"] table thead tr {
        text-align: center !important;
    }
    
    /* 最终兜底规则 - 使用最高权重 */
    .stDataFrame table tbody tr td,
    .stDataFrame table thead tr th {
        text-align: center !important;
        vertical-align: middle !important;
    }
    
    .stDataFrame table tbody tr td *,
    .stDataFrame table thead tr th * {
        text-align: center !important;
        vertical-align: middle !important;
    }
    /* 分栏间距优化 */
    .st-col {
        padding: 0 12px !important;
    }
    .dot-ani {
        display: inline-block;
        width: 1.2em;
        text-align: left;
    }
    .dot-ani:after {
        content: '...';
        animation: dots 1.2s steps(3, end) infinite;
    }
    @keyframes dots {
        0%, 20% { color: rgba(44,62,80,0); }
        40% { color: #2c3e50; }
        60% { color: #2c3e50; }
        80%, 100% { color: rgba(44,62,80,0); }
    }
</style>
""", unsafe_allow_html=True)

# 字体配置提示信息
def show_font_info():
    """显示字体配置信息"""
    import platform
    system = platform.system()
    
    with st.expander("🎨 字体配置说明", expanded=False):
        st.markdown("""
        ### 中文字体配置
        
        为了获得最佳的中文显示效果，建议安装思源黑体：
        
        **思源黑体下载链接：**
        - GitHub: https://github.com/adobe-fonts/source-han-sans
        - 直接下载：[SourceHanSansCN-Regular.otf](https://github.com/adobe-fonts/source-han-sans/releases)
        
        **安装路径建议：**
        """)
        
        if system == "Darwin":  # macOS
            st.code("""
macOS 安装路径：
- 系统级：/Library/Fonts/SourceHanSansCN-Regular.otf
- 用户级：~/Library/Fonts/SourceHanSansCN-Regular.otf

安装方法：下载字体文件后双击安装
            """, language="bash")
        elif system == "Linux":
            st.code("""
Linux 安装路径：
- 用户级：~/.local/share/fonts/SourceHanSansCN-Regular.otf
- 系统级：/usr/share/fonts/truetype/source-han-sans/

安装命令：
mkdir -p ~/.local/share/fonts
cp SourceHanSansCN-Regular.otf ~/.local/share/fonts/
fc-cache -fv
            """, language="bash")
        elif system == "Windows":
            st.code("""
Windows 安装路径：
- C:/Windows/Fonts/SourceHanSansCN-Regular.otf

安装方法：下载字体文件后右键选择"安装"
            """, language="bash")
        
        # 显示当前字体状态
        current_fonts = plt.rcParams.get('font.family', [])
        st.markdown(f"**当前字体配置：** `{', '.join(current_fonts)}`")
        
        if 'Source Han Sans CN' in current_fonts:
            st.success("✅ 思源黑体已成功配置")
        else:
            st.warning("⚠️ 未检测到思源黑体，使用系统默认中文字体")

# 主函数

def generate_word_report(df_filtered, total_cost, device_count, avg_monthly, dept_count, df):
    """生成Word分析报告（包含图表，按web界面顺序）"""
    doc = Document()
    import docx
    # 设置页边距为1厘米
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.3937)  # 1厘米=0.3937英寸
    
    # 设置文档默认字体
    def set_chinese_font_for_document(doc):
        """为Word文档设置中文字体"""
        from docx.oxml.ns import qn
        
        # 设置文档的默认字体
        doc.styles['Normal'].font.name = 'Source Han Sans CN'  # 优先使用思源黑体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Source Han Sans CN')
        
        # 如果思源黑体不可用，回退到其他中文字体
        try:
            # 尝试其他中文字体
            for font_name in ['Source Han Sans CN', '思源黑体', 'Microsoft YaHei', '微软雅黑', 'SimHei', '黑体']:
                doc.styles['Normal'].font.name = font_name
                doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                break
        except Exception as e:
            print(f"Word字体设置警告: {e}")
    
    set_chinese_font_for_document(doc)
    
    # 设置文档标题
    title = doc.add_heading('IT设备月度租赁分析系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加生成时间
    now_str = datetime.now().strftime('%Y年%m月%d日 %H:%M')
    time_para = doc.add_paragraph(f'报告生成时间: {now_str}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 若有AI分析，插入AI分析模块（富文本分段、加粗、字号、颜色等美化，尽量还原web展示）
    ai_result = st.session_state.get('ai_result', "")
    # 无论有无AI分析内容，始终插入AI分析区块（保持不隐藏）
    add_section_title(doc, 'AI智能分析')
    if ai_result:
        # 支持html富文本解析，100%还原web样式
        from bs4 import BeautifulSoup
        import re
        def parse_html_to_word(html, doc):
            soup = BeautifulSoup(html, 'html.parser')
            def add_run_with_style(paragraph, text, bold=False, color=None, size=13, font='Source Han Sans CN'):
                run = paragraph.add_run(text)
                run.font.name = font
                run.font.size = docx.shared.Pt(size)
                run.bold = bold
                if color:
                    run.font.color.rgb = color
                else:
                    run.font.color.rgb = RGBColor(44, 62, 80)
            def handle_tag(tag, parent=None):
                if tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    p = doc.add_paragraph()
                    size = 18 if tag.name in ['h1', 'h2'] else 16
                    add_run_with_style(p, tag.get_text(strip=True), bold=True, color=RGBColor(30, 64, 175), size=size)
                    p.paragraph_format.space_before = docx.shared.Pt(10)
                    p.paragraph_format.space_after = docx.shared.Pt(6)
                elif tag.name in ['b', 'strong']:
                    if parent:
                        add_run_with_style(parent, tag.get_text(), bold=True)
                elif tag.name == 'span':
                    style = tag.get('style', '')
                    color = None
                    size = 13
                    if 'color:' in style:
                        color_str = re.search(r'color:\s*([^;]+)', style)
                        if color_str:
                            color_val = color_str.group(1).strip()
                            if color_val.startswith('#'):
                                color = RGBColor(int(color_val[1:3],16), int(color_val[3:5],16), int(color_val[5:7],16))
                    if 'font-size:' in style:
                        size_str = re.search(r'font-size:\s*(\d+)px', style)
                        if size_str:
                            size = int(size_str.group(1)) * 0.75  # px转pt
                    if parent:
                        add_run_with_style(parent, tag.get_text(), size=size, color=color)
                elif tag.name in ['ul', 'ol']:
                    for li in tag.find_all('li', recursive=False):
                        p = doc.add_paragraph(style='List Bullet' if tag.name=='ul' else 'List Number')
                        add_run_with_style(p, li.get_text(strip=True))
                elif tag.name == 'li':
                    p = doc.add_paragraph(style='List Bullet')
                    add_run_with_style(p, tag.get_text(strip=True))
                elif tag.name == 'br':
                    doc.add_paragraph()
                elif tag.name == 'p' or tag.name is None:
                    p = doc.add_paragraph()
                    add_run_with_style(p, tag.get_text(strip=True))
                else:
                    # 递归处理未知标签
                    for child in tag.children:
                        if hasattr(child, 'name'):
                            handle_tag(child, parent)
                        else:
                            if parent:
                                add_run_with_style(parent, str(child))
            # 处理body下所有内容
            for elem in soup.contents:
                if hasattr(elem, 'name'):
                    handle_tag(elem)
                else:
                    p = doc.add_paragraph()
                    add_run_with_style(p, str(elem))
        # 判断ai_result是否为html
        if '<' in ai_result and '>' in ai_result:
            parse_html_to_word(ai_result, doc)
        else:
            # 兼容原有纯文本分段、加粗、标题、列表、缩进等
            paragraphs = [p.strip() for p in re.split(r'\n+', ai_result) if p.strip()]
            for para in paragraphs:
                # 标题识别
                heading_match = re.match(r'^[\s\u3000]*[\*·•-]?\s*([【\[]?)([\u4e00-\u9fa5A-Za-z0-9]+分析|总结|风险提示|建议|品牌分析|资产分析|费用分析|平台分析|状态分析|分类分析|人员分析|明细|结论|综述|整体分析|Top\d+)[】\]]?\s*$', para)
                bullet_match = re.match(r'^[\s\u3000]*([-•·*])\s+(.*)', para)
                indent_match = re.match(r'^(\s+)(.*)', para)
                run_bold = False
                color = None
                style = None
                left_indent = None
                text = para
                if heading_match:
                    heading_text = heading_match.group(2)
                    p = doc.add_paragraph()
                    run = p.add_run(heading_text)
                    run.font.size = docx.shared.Pt(16)
                    run.font.bold = True
                    run.font.name = 'Source Han Sans CN'
                    run.font.color.rgb = RGBColor(30, 64, 175)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = docx.shared.Pt(10)
                    p.paragraph_format.space_after = docx.shared.Pt(6)
                    continue
                if para.startswith('风险') or para.startswith('【风险'):
                    run_bold = True
                    color = RGBColor(220, 38, 38)
                elif para.startswith('建议') or para.startswith('【建议'):
                    run_bold = True
                    color = RGBColor(30, 64, 175)
                elif para.startswith('优化') or para.startswith('【优化'):
                    run_bold = True
                    color = RGBColor(16, 185, 129)
                if bullet_match:
                    text = bullet_match.group(2).strip()
                    style = 'List Bullet'
                elif indent_match:
                    left_indent = len(indent_match.group(1).replace('\u3000', '    ')) * 0.5
                    text = indent_match.group(2)
                p = doc.add_paragraph(style=style)
                run = p.add_run(text)
                run.font.size = docx.shared.Pt(15)
                run.font.name = 'Source Han Sans CN'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if left_indent:
                    p.paragraph_format.left_indent = docx.shared.Pt(left_indent * 12)
                if run_bold:
                    run.bold = True
                if color:
                    run.font.color.rgb = color
                if not color:
                    run.font.color.rgb = RGBColor(44, 62, 80)
                p.paragraph_format.space_before = docx.shared.Pt(2)
                p.paragraph_format.space_after = docx.shared.Pt(6)
            if not paragraphs:
                p = doc.add_paragraph()
                run = p.add_run(ai_result)
                run.font.size = docx.shared.Pt(15)
                run.font.name = 'Source Han Sans CN'
                run.font.color.rgb = RGBColor(44, 62, 80)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = docx.shared.Pt(2)
                p.paragraph_format.space_after = docx.shared.Pt(6)
    else:
        # 没有AI分析内容时，插入提示
        p = doc.add_paragraph()
        run = p.add_run('暂无AI分析内容。')
        run.font.size = docx.shared.Pt(15)
        run.font.name = 'Source Han Sans CN'
        run.font.color.rgb = RGBColor(180, 180, 180)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = docx.shared.Pt(2)
        p.paragraph_format.space_after = docx.shared.Pt(6)

    # 1. 关键指标概览（采用与下方表格一致的字体格式）
    add_section_title(doc, '1. 关键指标概览')
    
    # 创建4列表格来模拟web页面的4个卡片布局
    overview_table = doc.add_table(rows=2, cols=4)
    overview_table.autofit = True
    
    # 第一行：指标标题
    metrics = ['总租赁费用', '设备总数', '平均设备月租', '部门数量']
    for i, metric in enumerate(metrics):
        cell = overview_table.cell(0, i)
        cell.text = metric
        # 设置与beautify_table一致的标题样式
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = docx.shared.Pt(13)
                run.font.name = 'Source Han Sans CN'
                run.font.color.rgb = RGBColor(0, 0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置标题行背景色
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(r'<w:shd {} w:fill="D9EAFB"/>'.format(nsdecls('w'))))
    
    # 第二行：指标数值
    values = [f'¥{total_cost:,.2f}', f'{device_count}', f'¥{avg_monthly:,.2f}', f'{dept_count}']
    for i, value in enumerate(values):
        cell = overview_table.cell(1, i)
        cell.text = value
        # 设置与beautify_table一致的数值样式
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = docx.shared.Pt(13)
                run.font.name = 'Source Han Sans CN'
                run.font.color.rgb = RGBColor(0, 0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 设置表格边框（与beautify_table一致的黑色边框）
    for row in overview_table.rows:
        # 设置行高
        row.height = Cm(1.2)  # 设置行高为1.2厘米
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="8" w:color="000000"/><w:left w:val="single" w:sz="8" w:color="000000"/><w:bottom w:val="single" w:sz="8" w:color="000000"/><w:right w:val="single" w:sz="8" w:color="000000"/></w:tcBorders>'.format(nsdecls('w'))))

    # 2. 各平台关键指标分析（采用web页面样式的卡片布局）
    add_section_title(doc, '2. 各平台关键指标分析')
    
    # 设备数：排除包含'费'或'赔偿'字段的记录
    df_platform_device_filtered = filter_device_count_data(df_filtered)
    platform_device_data = df_platform_device_filtered.groupby('供应商')['实际金额'].count()
    
    # 总费用：使用所有数据
    platform_cost_data = df_filtered.groupby('供应商')['实际金额'].sum()
    
    # 合并数据
    platform_data = pd.DataFrame({
        '设备数': platform_device_data,
        '总费用': platform_cost_data
    }).fillna(0)
    
    # 平台顺序：易点云、小熊U租、其他
    supplier_order = ['易点云', '小熊U租'] + [s for s in platform_data.index if s not in ['易点云', '小熊U租']]
    platform_data = platform_data.reindex(supplier_order).dropna(how='all')
    
    if not platform_data.empty:
        # 创建2列表格来模拟web页面的并排卡片布局
        platform_cards_table = doc.add_table(rows=1, cols=2)
        platform_cards_table.autofit = True
        
        # 易点云卡片
        cell1 = platform_cards_table.cell(0, 0)
        if '易点云' in platform_data.index:
            yd_device_count = int(platform_data.loc['易点云', '设备数'])
            yd_total_cost = platform_data.loc['易点云', '总费用']
            apple_avg = df_filtered[(df_filtered['供应商']=='易点云') & (df_filtered['品牌类别']=='苹果')]['实际金额'].mean()
            win_avg = df_filtered[(df_filtered['供应商']=='易点云') & (df_filtered['品牌类别']=='Windows')]['实际金额'].mean()
            
            # 清空单元格内容
            cell1.text = ''
            
            # 添加标题
            title_p = cell1.add_paragraph()
            title_run = title_p.add_run('易点云')
            title_run.font.size = docx.shared.Pt(16)
            title_run.font.name = 'Source Han Sans CN'
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加指标数据
            metrics = [
                ('设备数', f'{yd_device_count}'),
                ('总费用', f'¥{yd_total_cost:,.2f}'),
                ('平均月租(苹果)', f'¥{apple_avg if not pd.isna(apple_avg) else 0:.2f}'),
                ('平均月租(Windows)', f'¥{win_avg if not pd.isna(win_avg) else 0:.2f}')
            ]
            
            for label, value in metrics:
                metric_p = cell1.add_paragraph()
                label_run = metric_p.add_run(f'{label}: ')
                label_run.font.size = docx.shared.Pt(13)
                label_run.font.name = 'Source Han Sans CN'
                label_run.font.bold = True
                label_run.font.color.rgb = RGBColor(0, 0, 0)
                
                value_run = metric_p.add_run(value)
                value_run.font.size = docx.shared.Pt(13)
                value_run.font.name = 'Source Han Sans CN'
                value_run.font.color.rgb = RGBColor(0, 0, 0)
                
                metric_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            cell1.text = '易点云\n无数据'
            # 设置无数据时的格式
            for p in cell1.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = docx.shared.Pt(13)
                    run.font.name = 'Source Han Sans CN'
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 小熊U租卡片
        cell2 = platform_cards_table.cell(0, 1)
        if '小熊U租' in platform_data.index:
            xz_device_count = int(platform_data.loc['小熊U租', '设备数'])
            xz_total_cost = platform_data.loc['小熊U租', '总费用']
            apple_avg = df_filtered[(df_filtered['供应商']=='小熊U租') & (df_filtered['品牌类别']=='苹果')]['实际金额'].mean()
            win_avg = df_filtered[(df_filtered['供应商']=='小熊U租') & (df_filtered['品牌类别']=='Windows')]['实际金额'].mean()
            
            # 清空单元格内容
            cell2.text = ''
            
            # 添加标题
            title_p = cell2.add_paragraph()
            title_run = title_p.add_run('小熊U租')
            title_run.font.size = docx.shared.Pt(16)
            title_run.font.name = 'Source Han Sans CN'
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加指标数据
            metrics = [
                ('设备数', f'{xz_device_count}'),
                ('总费用', f'¥{xz_total_cost:,.2f}'),
                ('平均月租(苹果)', f'¥{apple_avg if not pd.isna(apple_avg) else 0:.2f}'),
                ('平均月租(Windows)', f'¥{win_avg if not pd.isna(win_avg) else 0:.2f}')
            ]
            
            for label, value in metrics:
                metric_p = cell2.add_paragraph()
                label_run = metric_p.add_run(f'{label}: ')
                label_run.font.size = docx.shared.Pt(13)
                label_run.font.name = 'Source Han Sans CN'
                label_run.font.bold = True
                label_run.font.color.rgb = RGBColor(0, 0, 0)
                
                value_run = metric_p.add_run(value)
                value_run.font.size = docx.shared.Pt(13)
                value_run.font.name = 'Source Han Sans CN'
                value_run.font.color.rgb = RGBColor(0, 0, 0)
                
                metric_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            cell2.text = '小熊U租\n无数据'
            # 设置无数据时的格式
            for p in cell2.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = docx.shared.Pt(13)
                    run.font.name = 'Source Han Sans CN'
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 设置卡片样式（与beautify_table一致）
        for row in platform_cards_table.rows:
            # 设置行高
            row.height = Cm(1.2)  # 设置行高为1.2厘米
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 设置黑色边框
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="8" w:color="000000"/><w:left w:val="single" w:sz="8" w:color="000000"/><w:bottom w:val="single" w:sz="8" w:color="000000"/><w:right w:val="single" w:sz="8" w:color="000000"/></w:tcBorders>'.format(nsdecls('w'))))
                # 设置标题行背景色
                tcPr.append(parse_xml(r'<w:shd {} w:fill="D9EAFB"/>'.format(nsdecls('w'))))
        
        # 添加一些间距
        doc.add_paragraph()
        
        # 继续原有的图表和表格显示（保持现有功能）
        # 柱状图标题 - 美化显示
        add_chart_title(doc, '各平台总费用分布')
        
        # 柱状图（按固定顺序：易点云、小熊U租、其他）
        ensure_chinese_font()
        plt.figure(figsize=(10, 6))
        # 固定供应商顺序：易点云、小熊U租、其他
        supplier_order = ['易点云', '小熊U租'] + [s for s in platform_data.index if s not in ['易点云', '小熊U租']]
        platform_ordered = platform_data.reindex(supplier_order).dropna()
        
        # 为每个供应商分配不同颜色
        colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75']
        bar_colors = [colors[i % len(colors)] for i in range(len(platform_ordered))]
        
        bars = plt.bar(platform_ordered.index, platform_ordered['总费用'], color=bar_colors)
        plt.xlabel('供应商')
        plt.ylabel('总费用')
        plt.xticks(rotation=45)
        
        # 添加图例
        legend_elements = [plt.Rectangle((0,0),1,1, facecolor=bar_colors[i], label=supplier) 
                          for i, supplier in enumerate(platform_ordered.index)]
        plt.legend(handles=legend_elements, loc='upper right')
        
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height, f'¥{height:,.0f}', ha='center', va='bottom')
        plt.tight_layout()
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 饼图标题 - 美化显示
        add_chart_title(doc, '各平台费用占比')
        
        # 饼图
        ensure_chinese_font()
        plt.figure(figsize=(8, 6))
        colors = plt.cm.Set3.colors
        patches, texts, autotexts = plt.pie(platform_ordered['总费用'], labels=platform_ordered.index, autopct='%1.1f%%', startangle=90, colors=colors)
        plt.axis('equal')
        plt.legend(patches, platform_ordered.index, title="平台", bbox_to_anchor=(1, 0.5), loc="center left")
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 各平台关键指标明细表标题
        add_chart_title(doc, '各平台关键指标明细表')
        
        # 各平台关键指标明细表
        platform_detail_table = doc.add_table(rows=len(platform_ordered)+1, cols=5)
        platform_detail_table.autofit = True
        
        # 表格标题行
        headers = ['供应商', '设备数', '总费用', '平均月租(苹果)', '平均月租(Windows)']
        for col, header in enumerate(headers):
            platform_detail_table.cell(0, col).text = header
        
        # 表格数据行
        for i, (supplier, data) in enumerate(platform_ordered.iterrows()):
            # 计算苹果和Windows平均月租
            apple_avg = df_filtered[(df_filtered['供应商']==supplier) & (df_filtered['品牌类别']=='苹果')]['实际金额'].mean()
            win_avg = df_filtered[(df_filtered['供应商']==supplier) & (df_filtered['品牌类别']=='Windows')]['实际金额'].mean()
            
            # 供应商名称
            platform_detail_table.cell(i+1, 0).text = supplier
            # 设备数
            platform_detail_table.cell(i+1, 1).text = f'{int(data["设备数"])}'
            # 总费用
            platform_detail_table.cell(i+1, 2).text = f'¥{data["总费用"]:,.2f}'
            # 平均月租(苹果)
            platform_detail_table.cell(i+1, 3).text = f'¥{apple_avg if not pd.isna(apple_avg) else 0:.2f}'
            # 平均月租(Windows)
            platform_detail_table.cell(i+1, 4).text = f'¥{win_avg if not pd.isna(win_avg) else 0:.2f}'
        
        # 使用beautify_table函数统一设置表格样式
        beautify_table(platform_detail_table)
    
    # 3. 各平台设备数量（分品牌）分析（web顺序：柱状图-饼图-表格）
    add_section_title(doc, '3. 各平台设备数量（分品牌）分析')
    
    # 添加说明文字
    p = doc.add_paragraph()
    p.add_run('注：').bold = True
    p.add_run('品牌分析已排除品牌类别名称中包含"费"字的记录。')
    
    # 应用过滤函数排除包含'费'或'赔偿'字段的记录，同时排除品牌类别中包含'费'字的记录
    df_platform_filtered = filter_device_count_data(df_filtered)
    df_platform_filtered = df_platform_filtered[~df_platform_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
    platform_brand_devices = df_platform_filtered.groupby(['供应商', '品牌类别']).size().reset_index(name='设备数量')
    if not platform_brand_devices.empty:
        # 【颜色示例】可自定义分品牌柱状图配色，如：
        # colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75']
        # 用法：ax.bar(..., color=colors[i % len(colors)])
        # 柱状图标题 - 美化显示
        add_chart_title(doc, '各平台设备数量分布（按品牌）')
        
        # 柱状图（按固定顺序：易点云、小熊U租、其他）
        suppliers_unique = platform_brand_devices['供应商'].unique()
        supplier_order = ['易点云', '小熊U租'] + [s for s in suppliers_unique if s not in ['易点云', '小熊U租']]
        suppliers = [s for s in supplier_order if s in suppliers_unique]
        brands = platform_brand_devices['品牌类别'].unique()
        
        fig, ax = plt.subplots(figsize=(12, 7))  # 增加图表宽度
        x = np.arange(len(suppliers))
        
        # 根据品牌数量动态调整柱子宽度，避免重叠
        max_width = 0.8  # 所有柱子的最大总宽度
        width = max_width / len(brands) if len(brands) > 0 else 0.35
        width = min(width, 0.35)  # 限制最大宽度
        
        # 为每个品牌分配不同颜色
        colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75', '#FF6B6B', '#4ECDC4']
        
        for i, brand in enumerate(brands):
            brand_data = []
            for supplier in suppliers:
                count = platform_brand_devices[
                    (platform_brand_devices['供应商'] == supplier) & 
                    (platform_brand_devices['品牌类别'] == brand)
                ]['设备数量'].sum()
                brand_data.append(count)
            
            # 计算每个品牌柱子的x位置，确保不重叠
            x_pos = x + (i - len(brands)/2 + 0.5) * width
            color = colors[i % len(colors)]
            bars = ax.bar(x_pos, brand_data, width, label=brand, color=color, alpha=0.8)
            
            # 添加数值标签
            for j, bar in enumerate(bars):
                height = bar.get_height()
                if height > 0:
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{int(height)}台', ha='center', va='bottom', fontsize=9)
        
        ax.set_xlabel('供应商', fontsize=12)
        ax.set_ylabel('设备数量', fontsize=12)
        ax.set_xticks(x)
        ax.set_xticklabels(suppliers, rotation=0 if len(suppliers) <= 3 else 15)
        
        # 优化图例位置，避免与柱状图重叠
        if len(brands) <= 4:
            ax.legend(title="品牌类别", loc='upper right', fontsize=10)
        else:
            ax.legend(title="品牌类别", bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=9)
        
        # 设置图表边距，确保图例不被截断
        plt.subplots_adjust(right=0.85 if len(brands) > 4 else 0.95)
        plt.tight_layout()
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight',
                    facecolor='white', edgecolor='none')  # 设置背景色避免透明问题
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(7.5))  # 稍微增加宽度以适应更宽的图表
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 饼图标题 - 美化显示
        add_chart_title(doc, '各平台设备数量占比')
        
        # 饼图（平台总设备占比）
        platform_total = df_platform_filtered['供应商'].value_counts().reset_index()
        platform_total.columns = ['供应商', '设备数量']
        platform_total['供应商'] = pd.Categorical(platform_total['供应商'], categories=suppliers, ordered=True)
        platform_total = platform_total.sort_values('供应商')
        ensure_chinese_font()
        fig_platform_pie = plt.figure(figsize=(8, 6))
        plt.pie(platform_total['设备数量'], labels=platform_total['供应商'], autopct='%1.1f%%', startangle=90, colors=plt.cm.Set3.colors)
        plt.axis('equal')
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 明细表格标题
        add_chart_title(doc, '各平台设备数量明细表')
        # 明细表格
        brand_device_table = doc.add_table(rows=len(platform_brand_devices)+1, cols=3)
        brand_device_table.autofit = True
        headers = ['供应商', '品牌类别', '设备数量']
        for col, header in enumerate(headers):
            brand_device_table.cell(0, col).text = header
        for i, row in enumerate(platform_brand_devices.itertuples()):
            brand_device_table.cell(i+1, 0).text = row.供应商
            brand_device_table.cell(i+1, 1).text = row.品牌类别
            brand_device_table.cell(i+1, 2).text = f'{row.设备数量} 台'
        beautify_table(brand_device_table)
    
    # 4. 供应商平台分析（web顺序：柱状图-表格）
    add_section_title(doc, '4. 供应商平台分析')
    
    # 排除品牌类别中包含'费'字的记录
    df_brand_filtered = df_filtered[~df_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
    brand_detail = df_brand_filtered.groupby(['供应商', '品牌类别'])['实际金额'].agg(['sum', 'count', 'mean']).reset_index()
    brand_detail.columns = ['供应商', '品牌类别', '总价', '数量', '均价']
    if not brand_detail.empty:
        # 【颜色示例】可自定义平台-类别柱状图配色，如：
        # colors = ['#5B9BD5', '#ED7D31', '#A5A5A5', '#FFC000', '#4472C4', '#70AD47']
        # 用法：ax.bar(..., color=colors[i % len(colors)])
        # 柱状图标题
        add_chart_title(doc, '各供应商类别费用分布')
        # 柱状图（按固定顺序：易点云、小熊U租、其他）
        suppliers_unique = brand_detail['供应商'].unique()
        supplier_order = ['易点云', '小熊U租'] + [s for s in suppliers_unique if s not in ['易点云', '小熊U租']]
        suppliers = [s for s in supplier_order if s in suppliers_unique]
        brands = brand_detail['品牌类别'].unique()
        x = np.arange(len(suppliers))
        
        # 根据品牌数量动态调整柱子宽度，避免重叠
        max_width = 0.8  # 所有柱子的最大总宽度
        width = max_width / len(brands) if len(brands) > 0 else 0.35
        width = min(width, 0.35)  # 限制最大宽度，避免柱子过宽
        
        fig, ax = plt.subplots(figsize=(12, 7))  # 增加图表宽度以容纳更多柱子
        
        # 为每个品牌分配不同颜色
        colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75', '#FF6B6B', '#4ECDC4']
        
        for i, brand in enumerate(brands):
            brand_costs = []
            for supplier in suppliers:
                cost = brand_detail[
                    (brand_detail['供应商'] == supplier) & 
                    (brand_detail['品牌类别'] == brand)
                ]['总价'].sum()
                brand_costs.append(cost)
            
            # 计算每个品牌柱子的x位置，确保不重叠
            x_pos = x + (i - len(brands)/2 + 0.5) * width
            color = colors[i % len(colors)]
            bars = ax.bar(x_pos, brand_costs, width, label=brand, color=color, alpha=0.8)
            
            # 添加数值标签
            for j, bar in enumerate(bars):
                height = bar.get_height()
                if height > 0:
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'¥{height:,.0f}', ha='center', va='bottom', fontsize=9)
        
        ax.set_xlabel('供应商', fontsize=12)
        ax.set_ylabel('费用(元)', fontsize=12)
        ax.set_xticks(x)
        ax.set_xticklabels(suppliers, rotation=0 if len(suppliers) <= 3 else 15)  # 供应商多时适当旋转
        
        # 优化图例位置，避免与柱状图重叠
        if len(brands) <= 4:
            ax.legend(title="品牌类别", loc='upper right', fontsize=10)
        else:
            ax.legend(title="品牌类别", bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=9)
        
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'¥{x:,.0f}'))
        
        # 设置图表边距，确保图例不被截断
        plt.subplots_adjust(right=0.85 if len(brands) > 4 else 0.95)
        plt.tight_layout()
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight', 
                    facecolor='white', edgecolor='none')  # 设置背景色避免透明问题
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(7.5))  # 稍微增加宽度以适应更宽的图表
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 明细表格标题
        add_chart_title(doc, '供应商类别明细表')
        # 明细表格
        supplier_table = doc.add_table(rows=len(brand_detail)+1, cols=5)
        supplier_table.style = 'Table Grid'
        supplier_table.cell(0, 0).text = '供应商'
        supplier_table.cell(0, 1).text = '类别'
        supplier_table.cell(0, 2).text = '总价'
        supplier_table.cell(0, 3).text = '数量'
        supplier_table.cell(0, 4).text = '均价'
        for i, row in enumerate(brand_detail.itertuples()):
            supplier_table.cell(i+1, 0).text = row.供应商
            supplier_table.cell(i+1, 1).text = row.品牌类别
            supplier_table.cell(i+1, 2).text = f'¥{row.总价:,.2f}'
            supplier_table.cell(i+1, 3).text = f'{int(row.数量)} 台'
            supplier_table.cell(i+1, 4).text = f'¥{row.均价:,.2f}'
        beautify_table(supplier_table)
    
    # 5. 资产状态分析（图表在上，表格在下，顺序与web一致）
    if '资产状态' in df_filtered.columns:
        add_section_title(doc, '5. 资产状态分析')
        
        # 状态分类映射函数
        def categorize_status(status):
            status_str = str(status).strip()
            if '领用' in status_str:
                return '领用'
            elif '空闲' in status_str:
                return '空闲'
            elif '退租' in status_str or '处置' in status_str:
                return '已处置'
            elif '费' in status_str or '运费' in status_str or '赔偿' in status_str:
                return '费用'
            else:
                # 其他状态根据内容判断
                if any(keyword in status_str for keyword in ['归还', '返还', '报废', '损坏']):
                    return '已处置'
                else:
                    return '领用'  # 默认归类为领用
        
        # 计算费用总额（使用所有数据）
        status_sum = df_filtered.groupby('资产状态')['实际金额'].sum()
        status_sum_df = status_sum.reset_index()
        status_sum_df.columns = ['资产状态', '总金额']
        status_sum_df['状态分类'] = status_sum_df['资产状态'].apply(categorize_status)
        
        # 计算设备数量（排除费用类状态）
        df_status_device_filtered = df_filtered[~df_filtered['资产状态'].apply(lambda x: categorize_status(x) == '费用')]
        status_count = df_status_device_filtered.groupby('资产状态').size()
        status_count_df = status_count.reset_index()
        status_count_df.columns = ['资产状态', '设备数量']
        status_count_df['状态分类'] = status_count_df['资产状态'].apply(categorize_status)
        
        # 按状态分类汇总
        category_sum = status_sum_df.groupby('状态分类')['总金额'].sum()
        category_count = status_count_df.groupby('状态分类')['设备数量'].sum()
        
        # 按指定顺序：领用-空闲-已处置-费用
        order_list = ['领用', '空闲', '已处置', '费用']
        category_sum = category_sum.reindex(order_list, fill_value=0)
        category_count = category_count.reindex(order_list, fill_value=0)
        
        if not category_sum.empty:
            # 【颜色示例】可自定义资产状态柱状图/饼图配色，如：
            # bar_colors = ['#2E75B6', '#E46C0A', '#A9D18E', '#FFD966', '#C00000']
            # pie_colors = ['#2E75B6', '#E46C0A', '#A9D18E', '#FFD966', '#C00000']
            # 用法：plt.bar(..., color=bar_colors)，plt.pie(..., colors=pie_colors)
            
            # 说明文字
            p = doc.add_paragraph()
            p.add_run('注：').bold = True
            p.add_run('资产状态按领用、空闲、已处置、费用四类展示。设备数量统计不包含费用类状态，费用统计包含所有记录。')
            
            # 柱状图标题
            add_chart_title(doc, '资产状态分类设备数量分布')
            # 柱状图 - 只显示非费用类的设备数量
            ensure_chinese_font()
            plt.figure(figsize=(10, 6))
            
            # 过滤出非费用类的分类
            non_fee_categories = [cat for cat in order_list if cat != '费用' and category_count[cat] > 0]
            non_fee_counts = [category_count[cat] for cat in non_fee_categories]
            
            # 设置颜色
            color_map = {'领用': '#2E75B6', '空闲': '#9BBB59', '已处置': '#C0504D', '费用': '#E46C0A'}
            colors = [color_map[cat] for cat in non_fee_categories]
            
            bars = plt.bar(non_fee_categories, non_fee_counts, color=colors)
            plt.xlabel('资产状态分类')
            plt.ylabel('设备数量')
            # x轴标签加数量
            plt.xticks(ticks=range(len(non_fee_categories)), 
                      labels=[f"{cat}({int(cnt)})" for cat, cnt in zip(non_fee_categories, non_fee_counts)], 
                      rotation=0)
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}台', ha='center', va='bottom')
            # 添加图例
            legend_elements = [plt.Rectangle((0,0),1,1, facecolor=color_map[cat], label=cat) 
                              for cat in non_fee_categories]
            plt.legend(handles=legend_elements, title='状态分类', loc='upper right')
            plt.tight_layout()
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 饼图标题 - 按状态分类
            add_chart_title(doc, '资产状态分类占比')
            # 分类饼图
            ensure_chinese_font()
            plt.figure(figsize=(10, 8))
            
            # 只显示有设备数量的分类（排除费用类）
            pie_categories = []
            pie_counts = []
            pie_colors = []
            for cat in non_fee_categories:
                if category_count[cat] > 0:
                    pie_categories.append(cat)
                    pie_counts.append(category_count[cat])
                    pie_colors.append(color_map[cat])
            
            if pie_counts:
                pie_labels = [f"{cat}({int(cnt)})" for cat, cnt in zip(pie_categories, pie_counts)]
                patches, texts, autotexts = plt.pie(pie_counts, labels=pie_labels, autopct='%1.1f%%', 
                                                   startangle=90, colors=pie_colors)
                plt.axis('equal')
                # 添加图例
                plt.legend(patches, pie_categories, title="状态分类", bbox_to_anchor=(1, 0.5), loc="center left")
            
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 详细表格标题
            add_chart_title(doc, '资产状态分类汇总表')
            # 汇总表格
            summary_table = doc.add_table(rows=len(order_list)+1, cols=4)
            summary_table.autofit = True
            summary_table.cell(0, 0).text = '状态分类'
            summary_table.cell(0, 1).text = '设备数量'
            summary_table.cell(0, 2).text = '总金额'
            summary_table.cell(0, 3).text = '说明'
            
            for i, category in enumerate(order_list):
                summary_table.cell(i+1, 0).text = category
                if category == '费用':
                    summary_table.cell(i+1, 1).text = '-'
                    summary_table.cell(i+1, 3).text = '不统计设备数量'
                else:
                    device_count = category_count[category]
                    summary_table.cell(i+1, 1).text = f'{int(device_count)} 台'
                    summary_table.cell(i+1, 3).text = f'{int(device_count)}台设备'
                
                total_amount = category_sum[category]
                summary_table.cell(i+1, 2).text = f'¥{total_amount:,.2f}'
            
            beautify_table(summary_table)
    
    # 6. 资产分类分析（图表在上，表格在下，顺序与web一致）
    if '资产分类' in df_filtered.columns:
        add_section_title(doc, '6. 资产分类分析')
        
        # 添加说明文字
        p = doc.add_paragraph()
        p.add_run('注：').bold = True
        p.add_run('资产分类分析已排除资产分类名称中包含"费"字的记录。')
        
        # 排除资产分类中包含'费'字的记录
        df_asset_filtered = df_filtered[~df_filtered['资产分类'].astype(str).str.contains('费', na=False)]
        asset_type_analysis = df_asset_filtered.groupby('资产分类')['实际金额'].agg(['sum', 'count']).sort_values('sum', ascending=False)
        total_devices = len(df_asset_filtered)
        if not asset_type_analysis.empty:
            # 【颜色示例】可自定义资产分类柱状图/饼图配色，如：
            # bar_colors = ['#548235', '#A9D18E', '#F4B084', '#ED7D31', '#5B9BD5']
            # pie_colors = ['#548235', '#A9D18E', '#F4B084', '#ED7D31', '#5B9BD5']
            # 用法：plt.bar(..., color=bar_colors)，plt.pie(..., colors=pie_colors)
            # 柱状图标题
            add_chart_title(doc, '资产分类设备数量分布')
            # 柱状图
            ensure_chinese_font()
            plt.figure(figsize=(10, 6))
            bars = plt.bar(asset_type_analysis.index, asset_type_analysis['count'], color=plt.cm.Set2.colors)
            plt.xlabel('资产分类')
            plt.ylabel('设备数量')
            # x轴标签加数量
            plt.xticks(ticks=range(len(asset_type_analysis.index)), labels=[f"{name}({int(cnt)})" for name, cnt in zip(asset_type_analysis.index, asset_type_analysis['count'])], rotation=45)
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}台', ha='center', va='bottom')
            # 添加图例
            legend_elements = [plt.Rectangle((0,0),1,1, facecolor=plt.cm.Set2.colors[i % len(plt.cm.Set2.colors)], label=asset_type) 
                              for i, asset_type in enumerate(asset_type_analysis.index)]
            plt.legend(handles=legend_elements, title='资产分类', loc='upper right')
            plt.tight_layout()
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 饼图标题
            add_chart_title(doc, '资产分类设备数量占比')
            # 饼图
            ensure_chinese_font()
            plt.figure(figsize=(8, 6))
            pie_labels = [f"{name}({int(cnt)})" for name, cnt in zip(asset_type_analysis.index, asset_type_analysis['count'])]
            patches, texts, autotexts = plt.pie(asset_type_analysis['count'], labels=pie_labels, autopct='%1.1f%%', startangle=90, colors=plt.cm.Set3.colors)
            plt.axis('equal')
            # 添加图例
            plt.legend(patches, asset_type_analysis.index, title="资产分类", bbox_to_anchor=(1, 0.5), loc="center left")
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 明细表格标题
            add_chart_title(doc, '资产分类明细表')
            # 明细表格
            asset_type_table = doc.add_table(rows=len(asset_type_analysis)+1, cols=4)
            asset_type_table.autofit = True
            asset_type_table.cell(0, 0).text = '资产分类'
            asset_type_table.cell(0, 1).text = '总费用'
            asset_type_table.cell(0, 2).text = '设备数量'
            asset_type_table.cell(0, 3).text = '占比'
            for i, (asset_type, data) in enumerate(asset_type_analysis.iterrows()):
                asset_type_table.cell(i+1, 0).text = str(asset_type)
                asset_type_table.cell(i+1, 1).text = f'¥{data["sum"]:,.2f}'
                asset_type_table.cell(i+1, 2).text = f'{int(data["count"])} 台'
                asset_type_table.cell(i+1, 3).text = f'{data["count"]/total_devices*100:.1f}%'
            beautify_table(asset_type_table)
    
    # 7. 部门费用分析（图表在上，表格在下，顺序与web一致）
    add_section_title(doc, '7. 部门费用分析')
    dept_cost = df_filtered.groupby('一级部门')['实际金额'].sum().sort_values(ascending=False)
    if not dept_cost.empty:
        # 柱状图标题
        add_chart_title(doc, '各部门费用分布')
        # 柱状图
        ensure_chinese_font()
        plt.figure(figsize=(12, 6))
        bars = plt.bar(dept_cost.index, dept_cost.values, color=plt.cm.Set2.colors)
        plt.xlabel('部门')
        plt.ylabel('费用(元)')
        plt.xticks(rotation=45)
        plt.gca().yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'¥{x:,.0f}'))
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height, f'¥{height:,.0f}', ha='center', va='bottom')
        plt.tight_layout()
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 饼图标题
        add_chart_title(doc, '各部门费用占比')
        # 饼图
        ensure_chinese_font()
        plt.figure(figsize=(8, 6))
        plt.pie(dept_cost.values, labels=dept_cost.index, autopct='%1.1f%%', startangle=90, colors=plt.cm.Set3.colors)
        plt.axis('equal')
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        doc.add_paragraph()
        doc.add_picture(img_buffer, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 明细表格标题
        add_chart_title(doc, '部门费用排名表')
        # 明细表格
        dept_table = doc.add_table(rows=len(dept_cost)+1, cols=3)
        dept_table.autofit = True
        dept_table.cell(0, 0).text = '排名'
        dept_table.cell(0, 1).text = '部门'
        dept_table.cell(0, 2).text = '费用'
        for i, (dept, cost) in enumerate(dept_cost.items()):
            dept_table.cell(i+1, 0).text = str(i+1)
            dept_table.cell(i+1, 1).text = dept
            dept_table.cell(i+1, 2).text = f'¥{cost:,.2f}'
        beautify_table(dept_table)
    
    # 8. 人员费用分析（图表在上，表格在下，顺序与web一致）
    if '领用人' in df_filtered.columns and '人员编号' in df_filtered.columns:
        add_section_title(doc, '8. 人员费用分析')
        # 应用过滤函数排除包含'费'或'赔偿'字段的记录
        df_person_filtered = filter_device_count_data(df_filtered)
        person_group = df_person_filtered.groupby(['领用人', '人员编号', '一级部门']).agg({
            '实际金额': 'sum',
            '品牌类别': lambda x: ', '.join(sorted(set(map(str, x)))) if len(x) else '',
            '领用人': 'count'
        }).rename(columns={'实际金额': '总费用', '品牌类别': '设备类型', '领用人': '设备数量'}).reset_index()
        # 只显示费用超过500元的人员
        person_group_filtered = person_group[person_group['总费用'] > 500].sort_values('总费用', ascending=False)
        if not person_group_filtered.empty:
            # 柱状图标题
            add_chart_title(doc, '人员费用分析(Top 10)')
            # 柱状图
            top10_persons = person_group_filtered.head(10)
            person_labels = [f"{row['领用人']}({row['人员编号']})" for _, row in top10_persons.iterrows()]
            ensure_chinese_font()
            plt.figure(figsize=(12, 6))
            bars = plt.bar(person_labels, top10_persons['总费用'], color=plt.cm.Set2.colors)
            plt.xlabel('人员')
            plt.ylabel('费用(元)')
            plt.xticks(rotation=45, ha='right')
            plt.gca().yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'¥{x:,.0f}'))
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height, f'¥{height:,.0f}', ha='center', va='bottom')
            plt.tight_layout()
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 饼图标题
            add_chart_title(doc, '人员费用占比(Top 10)')
            # 饼图
            ensure_chinese_font()
            plt.figure(figsize=(8, 6))
            plt.pie(top10_persons['总费用'], labels=person_labels, autopct='%1.1f%%', startangle=90, colors=plt.cm.Set3.colors)
            plt.axis('equal')
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 明细表格标题
            add_chart_title(doc, '人员费用排名表(Top 15)')
            # 明细表格（显示前15名）
            person_table = doc.add_table(rows=min(len(person_group_filtered), 15)+1, cols=5)
            person_table.autofit = True
            person_table.cell(0, 0).text = '排名'
            person_table.cell(0, 1).text = '姓名(工号)'
            person_table.cell(0, 2).text = '部门'
            person_table.cell(0, 3).text = '设备数量'
            person_table.cell(0, 4).text = '总费用'
            for i, row in enumerate(person_group_filtered.head(15).itertuples()):
                person_table.cell(i+1, 0).text = str(i+1)
                person_table.cell(i+1, 1).text = f'{row.领用人}({row.人员编号})'
                person_table.cell(i+1, 2).text = row.一级部门
                person_table.cell(i+1, 3).text = f'{row.设备数量} 台'
                person_table.cell(i+1, 4).text = f'¥{row.总费用:,.2f}'
            beautify_table(person_table)
    
    # 9. 人员领取多台设备分析（图表在上，表格在下，顺序与web一致）
    if '领用人' in df_filtered.columns and '人员编号' in df_filtered.columns:
        add_section_title(doc, '9. 人员领取多台设备分析')
        # 应用过滤函数排除包含'费'或'赔偿'字段的记录
        df_multi_device_filtered = filter_device_count_data(df_filtered)
        multi_device_all = df_multi_device_filtered.groupby(['领用人', '人员编号', '一级部门']).agg({
            '实际金额': 'sum',
            '品牌类别': lambda x: ', '.join(sorted(set(map(str, x)))) if len(x) else '',
            '领用人': 'count'
        }).rename(columns={'实际金额': '总费用', '品牌类别': '设备类型', '领用人': '设备数量'}).reset_index()
        multi_device_all = multi_device_all[multi_device_all['设备数量'] > 1]
        multi_device_all = multi_device_all.sort_values('设备数量', ascending=False)
        multi_device_all['人员'] = multi_device_all['领用人'].astype(str) + '（' + multi_device_all['人员编号'].astype(str) + '）'
        if not multi_device_all.empty:
            # 柱状图标题
            add_chart_title(doc, '人员多台设备领取情况')
            # 柱状图
            ensure_chinese_font()
            plt.figure(figsize=(12, 6))
            bars = plt.bar(multi_device_all['人员'], multi_device_all['设备数量'], color=plt.cm.Set2.colors)
            plt.xlabel('人员')
            plt.ylabel('设备数量')
            plt.xticks(rotation=45, ha='right')
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}台', ha='center', va='bottom')
            plt.tight_layout()
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close()
            doc.add_paragraph()
            doc.add_picture(img_buffer, width=Inches(8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 明细表标题
            add_chart_title(doc, '人员多台设备领取明细表')
            # 明细表
            table = doc.add_table(rows=len(multi_device_all)+1, cols=5)
            table.autofit = True
            headers = ['人员', '一级部门', '设备数量', '设备类型', '总费用']
            for col, header in enumerate(headers):
                table.cell(0, col).text = header
            for i, row in enumerate(multi_device_all.itertuples()):
                table.cell(i+1, 0).text = row.人员
                table.cell(i+1, 1).text = row.一级部门
                table.cell(i+1, 2).text = str(row.设备数量)
                table.cell(i+1, 3).text = row.设备类型
                table.cell(i+1, 4).text = f'¥{row.总费用:,.2f}'
            beautify_table(table)

    # 10. 分析总结
    add_section_title(doc, '10. 分析总结')
    summary_para = doc.add_paragraph()
    summary_para.add_run('本报告基于当前筛选条件生成，包含了设备租赁的关键指标分析和可视化图表。')
    summary_para.add_run('建议定期查看各部门设备使用情况，优化设备配置和成本控制。')
    # 添加数据说明
    doc.add_paragraph()
    data_para = doc.add_paragraph()
    data_para.add_run(f'数据范围：共分析 {device_count} 台设备，涉及 {dept_count} 个部门，总费用 ¥{total_cost:,.2f}。')

    # 保存到内存
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer

def generate_pdf_report(df_filtered, total_cost, device_count, avg_monthly, dept_count):
    """生成PDF分析报告（与web页面1:1还原）"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=60, bottomMargin=40)
    
    # 注册中文字体（优先思源黑体）
    def register_chinese_fonts():
        font_paths = [
            '/home/caosaikang/.local/share/fonts/SourceHanSansCN-Regular.otf',
            '/System/Library/Fonts/PingFang.ttc',  # macOS
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Linux
            'C:/Windows/Fonts/msyh.ttc',  # Windows
        ]
        
        registered_font = None
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if 'SourceHan' in font_path:
                        pdfmetrics.registerFont(TTFont('SourceHanSans', font_path))
                        registered_font = 'SourceHanSans'
                        print(f"PDF: 成功注册思源黑体 {font_path}")
                        break
                    elif 'PingFang' in font_path:
                        pdfmetrics.registerFont(TTFont('PingFang', font_path))
                        registered_font = 'PingFang'
                        break
                    elif 'msyh' in font_path:
                        pdfmetrics.registerFont(TTFont('MicrosoftYaHei', font_path))
                        registered_font = 'MicrosoftYaHei'
                        break
                except Exception as e:
                    print(f"注册字体失败 {font_path}: {e}")
                    continue
        
        return registered_font or 'Helvetica'
    
    chinese_font = register_chinese_fonts()
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 自定义样式（使用中文字体）
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,  # 居中
        textColor=colors.HexColor('#2c3e50'),
        fontName=chinese_font
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=20,
        spaceBefore=20,
        textColor=colors.HexColor('#2c3e50'),
        fontName=chinese_font
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        textColor=colors.HexColor('#333333'),
        fontName=chinese_font
    )
    
    # 创建故事内容
    story = []
    
    # 1. 报告标题
    story.append(Paragraph('💻 IT设备月度租赁分析系统', title_style))
    story.append(Paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}', normal_style))
    story.append(Spacer(1, 20))
    
    # 2. 数据摘要
    story.append(Paragraph('📊 数据摘要', heading_style))
    
    # 创建关键指标表格
    summary_data = [
        ['指标', '数值'],
        ['总租赁费用', f'¥{total_cost:,.2f}'],
        ['设备总数', f'{device_count} 台'],
        ['平均设备月租', f'¥{avg_monthly:,.2f}'],
        ['涉及部门数', f'{dept_count} 个']
    ]
    
    summary_table = Table(summary_data, colWidths=[2.5*inch, 2.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8f9fa')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('FONTSIZE', (0, 1), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e9ecef'))
    ]))
    
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # 3. 部门费用分析
    story.append(Paragraph('🏢 部门费用分析', heading_style))
    
    dept_cost = df_filtered.groupby('一级部门')['实际金额'].sum().sort_values(ascending=False)
    
    if not dept_cost.empty:
        # 部门费用表格
        dept_table_data = [['排名', '部门', '费用', '占比']]
        total_cost_dept = dept_cost.sum()
        
        for i, (dept, cost) in enumerate(dept_cost.items()):
            dept_table_data.append([
                str(i+1),
                dept,
                f'¥{cost:,.2f}',
                f'{cost/total_cost_dept*100:.1f}%'
            ])
        
        dept_table = Table(dept_table_data, colWidths=[1*inch, 2.5*inch, 1.5*inch, 1*inch])
        dept_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8f9fa')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e9ecef'))
        ]))
        
        story.append(dept_table)
        story.append(Spacer(1, 15))
        
        # 部门费用图表（使用plotly转图片）
        try:
            fig_dept_bar = px.bar(
                dept_cost.reset_index(),
                x='一级部门',
                y='实际金额',
                title='部门费用分布',
                text='实际金额'
            )
            fig_dept_bar.update_traces(texttemplate='¥%{text:,.0f}', textposition='outside')
            fig_dept_bar.update_layout(
                yaxis_title='费用(元)',
                xaxis_title='部门',
                font=dict(size=12),
                title_font_size=14,
                showlegend=False,
                width=800,
                height=400
            )
            
            img_bytes = fig_dept_bar.to_image(format="png", width=800, height=400, scale=2)
            img_buffer = io.BytesIO(img_bytes)
            chart_img = Image(img_buffer, width=6.5*inch, height=3.4*inch)
            story.append(chart_img)
            story.append(Spacer(1, 20))
        except Exception as e:
            story.append(Paragraph(f'图表生成失败: {str(e)}', normal_style))
            story.append(Spacer(1, 20))
    
    # 4. 品牌分析
    if '品牌类别' in df_filtered.columns:
        story.append(Paragraph('🏷️ 品牌分析', heading_style))
        
        # 排除品牌类别中包含'费'字的记录
        df_brand_pdf_filtered = df_filtered[~df_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
        brand_cost = df_brand_pdf_filtered.groupby('品牌类别')['实际金额'].sum().sort_values(ascending=False)
        brand_count = df_brand_pdf_filtered['品牌类别'].value_counts()
        
        if not brand_cost.empty:
            # 类别费用表格
            brand_table_data = [['排名', '类别', '设备数量', '总费用', '平均费用']]
            
            for i, brand in enumerate(brand_cost.index[:10]):  # 只显示前10
                cost = brand_cost[brand]
                count = brand_count[brand]
                avg_cost = cost / count if count else 0
                brand_table_data.append([
                    str(i+1),
                    brand,
                    f'{count} 台',
                    f'¥{cost:,.2f}',
                    f'¥{avg_cost:,.2f}'
                ])
            
            brand_table = Table(brand_table_data, colWidths=[0.8*inch, 2*inch, 1.2*inch, 1.5*inch, 1.5*inch])
            brand_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8f9fa')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e9ecef'))
            ]))
            
            story.append(brand_table)
            story.append(Spacer(1, 20))
    
    # 5. 供应商分析
    story.append(Paragraph('🏭 供应商分析', heading_style))
    
    supplier_cost = df_filtered.groupby('供应商')['实际金额'].sum().sort_values(ascending=False)
    supplier_count = df_filtered['供应商'].value_counts()
    
    if not supplier_cost.empty:
        # 供应商费用表格
        supplier_table_data = [['排名', '供应商', '设备数量', '总费用', '占比']]
        total_supplier_cost = supplier_cost.sum()
        
        for i, supplier in enumerate(supplier_cost.index[:10]):  # 只显示前10
            cost = supplier_cost[supplier]
            count = supplier_count[supplier]
            percentage = cost / total_supplier_cost * 100 if total_supplier_cost else 0
            supplier_table_data.append([
                str(i+1),
                supplier,
                f'{count} 台',
                f'¥{cost:,.2f}',
                f'{percentage:.1f}%'
            ])
        
        supplier_table = Table(supplier_table_data, colWidths=[0.8*inch, 2.2*inch, 1.2*inch, 1.5*inch, 1.3*inch])
        supplier_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8f9fa')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e9ecef'))
        ]))
        
        story.append(supplier_table)
        story.append(Spacer(1, 20))
    
    # 6. 分析总结
    story.append(Paragraph('📝 分析总结', heading_style))
    story.append(Paragraph('本报告基于当前筛选条件生成，包含了设备租赁的关键指标分析和可视化图表。建议定期查看各部门设备使用情况，优化设备配置和成本控制。', normal_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f'数据范围：共分析 {device_count} 台设备，涉及 {dept_count} 个部门，总费用 ¥{total_cost:,.2f}。', normal_style))
    
    # 构建PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def main():
    filter_disabled = False  # 保证所有用到它的地方都已定义，避免UnboundLocalError

    st.markdown('<h1 class="page-title">💻 IT设备月度租赁分析系统</h1>', unsafe_allow_html=True)
    st.markdown("数据分析 | 支持Excel账单上传 | 版本 2.4")

    # 上传Excel文件
    uploaded_file = st.file_uploader("上传设备账单Excel文件（支持xlsx/xls）", type=["xlsx", "xls"])
    if uploaded_file is None:
        st.info("请上传Excel账单文件以开始分析。")
        return
    try:
        df = pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"文件读取失败: {e}")
        return

    # 只保留“使用部门”字段的一级部门（遇到/分隔时取第一个）
    if "使用部门" in df.columns:
        df["一级部门"] = df["使用部门"].astype(str).str.split("/").str[0]
    else:
        st.error("上传的Excel缺少‘使用部门’字段，请检查模板！")
        return

    # 检查其它关键字段是否存在
    required_columns = ["品牌", "供应商", "应付金额", "所在位置"]
    missing_cols = [col for col in required_columns if col not in df.columns]
    if missing_cols:
        st.error(f"上传的Excel缺少关键字段: {', '.join(missing_cols)}，请检查模板！")
        return
    # 检查关键字段每行均有数据
    if df[required_columns].isnull().any(axis=1).any() or (df[required_columns] == '').any(axis=1).any():
        st.error("上传的Excel存在关键字段（品牌、供应商、应付金额、所在位置）缺失值，请补全后再上传！")
        return

    # 字段兼容性处理
    df.rename(columns={"品牌": "品牌类别", "应付金额": "实际金额"}, inplace=True)
    # 兼容“使用人”字段为“领用人”
    if "使用人" in df.columns and "领用人" not in df.columns:
        df.rename(columns={"使用人": "领用人"}, inplace=True)

    # 品牌类别标准化：包含“苹果”即为苹果，其次佳能=佳能，其余均为Windows
    def map_brand(x):
        x_str = str(x).strip().upper()
        if '苹果' in x_str or 'APPLE' in x_str:
            return '苹果'
        elif '佳能' in x_str or 'CANON' in x_str:
            return '佳能'
        else:
            return 'Windows'
    df['品牌类别'] = df['品牌类别'].apply(map_brand)

    # 检查“所在位置”字段是否存在
    if "所在位置" not in df.columns:
        st.error("上传的Excel缺少‘所在位置’字段，请检查模板！")
        return

    # 资产状态标准化：所有包含“处置”的状态统一为“退租”
    if '资产状态' in df.columns:
        df['资产状态'] = df['资产状态'].astype(str).apply(lambda x: '退租' if '处置' in x else x)

    # 控制筛选控件是否禁用（AI分析时锁定）
    filter_disabled = st.session_state.get('ai_analyzing', False)
    # 数据筛选区域（对应PDF中的筛选栏）
    with st.container():
        st.markdown('<h3 class="section-title">🔍 数据筛选</h3>', unsafe_allow_html=True)
        # --- 新增：清除筛选按钮 ---
        clear_col, _ = st.columns([1, 5])
        with clear_col:
            if st.button("清除筛选", key="clear_filter_btn", help="一键重置所有筛选条件为全部", disabled=filter_disabled):
                st.session_state['brand'] = "全部"
                st.session_state['dept'] = "全部"
                st.session_state['supplier'] = "全部"
                st.session_state['location'] = "全部"
                st.session_state['asset_status'] = "全部"
                st.session_state['asset_type'] = "全部"
                st.rerun()
        # --- 筛选控件，全部用session_state管理，并排排列 ---
        filter_cols = st.columns(6, gap="small")
        with filter_cols[0]:
            brand_options = ["全部"] + sorted(df['品牌类别'].unique().tolist())
            if 'brand' in st.session_state and st.session_state['brand'] in brand_options:
                brand_index = brand_options.index(st.session_state['brand'])
            else:
                brand_index = 0
            brand = st.selectbox("品牌类别", brand_options, index=brand_index, key="brand", disabled=filter_disabled)
        with filter_cols[1]:
            dept_options = ["全部"] + sorted(df['一级部门'].unique().tolist())
            if 'dept' in st.session_state and st.session_state['dept'] in dept_options:
                dept_index = dept_options.index(st.session_state['dept'])
            else:
                dept_index = 0
            dept = st.selectbox("部门", dept_options, index=dept_index, key="dept", disabled=filter_disabled)
        with filter_cols[2]:
            supplier_options = ["全部"] + sorted(df['供应商'].unique().tolist())
            if 'supplier' in st.session_state and st.session_state['supplier'] in supplier_options:
                supplier_index = supplier_options.index(st.session_state['supplier'])
            else:
                supplier_index = 0
            supplier = st.selectbox("供应商", supplier_options, index=supplier_index, key="supplier", disabled=filter_disabled)
        with filter_cols[3]:
            location_options = ["全部"] + sorted(df['所在位置'].unique().tolist())
            if 'location' in st.session_state and st.session_state['location'] in location_options:
                location_index = location_options.index(st.session_state['location'])
            else:
                location_index = 0
            location = st.selectbox("所在位置", location_options, index=location_index, key="location", disabled=filter_disabled)
        with filter_cols[4]:
            if '资产状态' in df.columns:
                asset_status_options = ["全部"] + sorted(df['资产状态'].unique().tolist())
            else:
                asset_status_options = ["全部"]
            if 'asset_status' in st.session_state and st.session_state['asset_status'] in asset_status_options:
                asset_status_index = asset_status_options.index(st.session_state['asset_status'])
            else:
                asset_status_index = 0
            asset_status = st.selectbox("资产状态", asset_status_options, index=asset_status_index, key="asset_status", disabled=filter_disabled)
        with filter_cols[5]:
            if '资产分类' in df.columns:
                asset_type_options = ["全部"] + sorted(df['资产分类'].unique().tolist())
            else:
                asset_type_options = ["全部"]
            if 'asset_type' in st.session_state and st.session_state['asset_type'] in asset_type_options:
                asset_type_index = asset_type_options.index(st.session_state['asset_type'])
            else:
                asset_type_index = 0
            asset_type = st.selectbox("资产分类", asset_type_options, index=asset_type_index, key="asset_type", disabled=filter_disabled)

        # 多条件筛选
        df_filtered = df.copy()
        if brand != "全部":
            df_filtered = df_filtered[df_filtered['品牌类别'] == brand]
        if dept != "全部":
            df_filtered = df_filtered[df_filtered['一级部门'] == dept]
        if supplier != "全部":
            df_filtered = df_filtered[df_filtered['供应商'] == supplier]
        if location != "全部":
            df_filtered = df_filtered[df_filtered['所在位置'] == location]
        if asset_status != "全部" and '资产状态' in df_filtered.columns:
            df_filtered = df_filtered[df_filtered['资产状态'] == asset_status]
        if asset_type != "全部" and '资产分类' in df_filtered.columns:
            df_filtered = df_filtered[df_filtered['资产分类'] == asset_type]

        # 显示筛选结果（修复分隔符和空项问题）
        filter_items = []
        if brand != "全部":
            filter_items.append(f"品牌类别: {brand}")
        if dept != "全部":
            filter_items.append(f"部门: {dept}")
        if supplier != "全部":
            filter_items.append(f"供应商: {supplier}")
        if location != "全部":
            filter_items.append(f"所在位置: {location}")
        if asset_status != "全部" and asset_status != "":
            filter_items.append(f"资产状态: {asset_status}")
        if asset_type != "全部" and asset_type != "":
            filter_items.append(f"资产分类: {asset_type}")
        if filter_items:
            filter_text = " | ".join(filter_items)
        else:
            filter_text = "全部"
        st.markdown(f"<p style='font-size: 14px; color: #6c757d; margin-top: 10px;'>当前筛选: {filter_text}</p>", unsafe_allow_html=True)
        # 显示设备数量时也应用过滤函数
        df_display_filtered = filter_device_count_data(df_filtered)
        st.markdown(f"<p style='font-size: 14px; color: #2c3e50; font-weight: 500;'>显示设备数: {len(df_display_filtered)} 台</p>", unsafe_allow_html=True)
    
    # 关键指标变量提前定义，供导出和分析区块复用
    total_cost = df_filtered['实际金额'].sum()
    # 应用过滤函数排除包含'费'或'赔偿'字段的记录
    df_device_count_filtered = filter_device_count_data(df_filtered)
    device_count = len(df_device_count_filtered)
    avg_monthly = total_cost / device_count if device_count else 0
    dept_count = len(df_filtered['一级部门'].unique())
    brand_count = len(df_filtered['品牌类别'].unique()) if '品牌类别' in df_filtered.columns else 0
    supplier_count = len(df_filtered['供应商'].unique()) if '供应商' in df_filtered.columns else 0
    asset_status_count = len(df_filtered['资产状态'].unique()) if '资产状态' in df_filtered.columns else 0
    asset_type_count = len(df_filtered['资产分类'].unique()) if '资产分类' in df_filtered.columns else 0
    person_count = len(df_filtered['人员编号'].unique()) if '人员编号' in df_filtered.columns else 0
    
    # 导出功能区域（放在筛选下方，靠左显示）
    st.markdown("---")
    st.markdown("### 📄 报告导出")
    
    # Word导出按钮
    if st.button("📝 生成Word分析报告", key="export_word_btn", help="生成当前筛选数据的Word分析报告"):
        try:
            with st.spinner('正在生成Word报告...'):
                # 生成Word文档，传入全量df
                word_buffer = generate_word_report(df_filtered, total_cost, device_count, avg_monthly, dept_count, df)
                # 生成文件名
                now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'IT设备租赁分析报告_{now_str}.docx'
                # 使用session_state存储生成的文档
                st.session_state['word_buffer'] = word_buffer
                st.session_state['word_filename'] = filename
                st.success("Word报告生成成功！")
        except Exception as e:
            st.error(f"生成Word报告失败：{e}")
    
    # 显示下载按钮（如果已生成）
    if 'word_buffer' in st.session_state and 'word_filename' in st.session_state:
        col1, col2 = st.columns([1, 1])
        with col1:
            st.download_button(
                label="⬇️ 下载Word报告",
                data=st.session_state['word_buffer'].getvalue(),
                file_name=st.session_state['word_filename'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word_btn"
            )
        with col2:
            if st.button("🔄 重新生成报告", key="regenerate_word_btn"):
                # 清除之前的生成结果
                if 'word_buffer' in st.session_state:
                    del st.session_state['word_buffer']
                if 'word_filename' in st.session_state:
                    del st.session_state['word_filename']
                st.rerun()
    

    # --- 全局AI整体分析区块，放在关键指标概览上方 ---
    with st.container():
        st.markdown('<h3 class="section-title">🤖 AI整体智能分析</h3>', unsafe_allow_html=True)
        # 构建 summary_text，供AI分析用
        summary_text = ""
        if '资产分类' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录，以及资产分类中包含'费'字的记录
            df_ai_filtered = filter_device_count_data(df_filtered)
            df_ai_filtered = df_ai_filtered[~df_ai_filtered['资产分类'].astype(str).str.contains('费', na=False)]
            if not df_ai_filtered.empty:
                top_type = df_ai_filtered['资产分类'].value_counts().idxmax()
                top_type_count = df_ai_filtered['资产分类'].value_counts().max()
                summary_text += f"设备数量最多资产分类: {top_type}（{top_type_count}台）。\n"
        # 平台分布
        # 应用过滤函数排除包含'费'或'赔偿'字段的记录
        df_platform_ai_filtered = filter_device_count_data(df_filtered)
        platform_data = df_platform_ai_filtered.groupby('供应商')['实际金额'].agg(['sum', 'count'])
        platform_summary = ', '.join([f"{idx}: ¥{row['sum']:,.2f}({int(row['count'])}台)" for idx, row in platform_data.iterrows()])
        summary_text += f"平台分布: {platform_summary}\n"
        # 部门分布
        dept_cost = df_filtered.groupby('一级部门')['实际金额'].sum().sort_values(ascending=False)
        dept_summary = ', '.join([f"{idx}: ¥{val:,.2f}" for idx, val in dept_cost.head(5).items()])
        summary_text += f"部门费用Top5: {dept_summary}\n"
        # 品牌分布
        if '品牌类别' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录，同时排除品牌类别中包含'费'字的记录
            df_brand_ai_filtered = filter_device_count_data(df_filtered)
            df_brand_ai_filtered = df_brand_ai_filtered[~df_brand_ai_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
            if not df_brand_ai_filtered.empty:
                brand_summary = ', '.join([f"{b}: {n}台" for b, n in df_brand_ai_filtered['品牌类别'].value_counts().items()])
                summary_text += f"品牌分布: {brand_summary}\n"
        # 资产状态
        if '资产状态' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录
            df_status_ai_filtered = filter_device_count_data(df_filtered)
            status_count = df_status_ai_filtered['资产状态'].value_counts()
            status_summary = ', '.join([f"{idx}: {val}台" for idx, val in status_count.items()])
            summary_text += f"资产状态分布: {status_summary}\n"
        # 资产分类
        if '资产分类' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录
            df_asset_type_ai_filtered = filter_device_count_data(df_filtered)
            asset_type_count_series = df_asset_type_ai_filtered['资产分类'].astype(str).value_counts()
            asset_type_summary = ', '.join([f"{idx}: {val}台" for idx, val in asset_type_count_series.items()])
            summary_text += f"资产分类分布: {asset_type_summary}\n"
        # 人员费用Top5
        if '领用人' in df_filtered.columns and '人员编号' in df_filtered.columns:
            person_group = df_filtered.groupby(['领用人', '人员编号']).agg({'实际金额': 'sum'}).reset_index()
            person_group = person_group.sort_values('实际金额', ascending=False).head(5)
            person_summary = ', '.join([f"{row['领用人']}({row['人员编号']}): ¥{row['实际金额']:,.2f}" for _, row in person_group.iterrows()])
            summary_text += f"人员费用Top5: {person_summary}\n"
        # 平台-类别费用分布
        if '供应商' in df_filtered.columns and '品牌类别' in df_filtered.columns:
            # 排除品牌类别中包含'费'字的记录
            df_platform_brand_ai_filtered = df_filtered[~df_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
            brand_platform_cost = df_platform_brand_ai_filtered.groupby(['供应商', '品牌类别'])['实际金额'].sum().reset_index()
            brand_platform_summary = ', '.join([f"{row['供应商']} - {row['品牌类别']}: ¥{row['实际金额']:,.2f}" for _, row in brand_platform_cost.iterrows()])
            summary_text += f"平台-类别费用分布: {brand_platform_summary}\n"
        # 设备单价分布
        if '实际金额' in df_filtered.columns:
            min_price = df_filtered['实际金额'].min()
            max_price = df_filtered['实际金额'].max()
            median_price = df_filtered['实际金额'].median()
            summary_text += f"设备单价区间: ¥{min_price:,.2f} ~ ¥{max_price:,.2f}，中位数: ¥{median_price:,.2f}\n"
        # 设备领取情况
        if '领用人' in df_filtered.columns and '人员编号' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录
            df_multi_ai_filtered = filter_device_count_data(df_filtered)
            multi_device = df_multi_ai_filtered.groupby(['领用人', '人员编号']).size().reset_index(name='设备数')
            multi_count = (multi_device['设备数'] > 1).sum()
            summary_text += f"领取多台设备人员数: {multi_count} 人\n"
        # 设备分布地
        if '所在位置' in df_filtered.columns:
            loc_summary = ', '.join([f"{loc}: {n}台" for loc, n in df_filtered['所在位置'].value_counts().items()])
            summary_text += f"设备分布地: {loc_summary}\n"
        # 资产分类金额Top3
        if '资产分类' in df_filtered.columns:
            # 排除资产分类中包含'费'字的记录
            df_asset_ai_filtered = df_filtered[~df_filtered['资产分类'].astype(str).str.contains('费', na=False)]
            asset_type_sum = df_asset_ai_filtered.groupby('资产分类')['实际金额'].sum().sort_values(ascending=False).head(3)
            asset_type_sum_summary = ', '.join([f"{idx}: ¥{val:,.2f}" for idx, val in asset_type_sum.items()])
            summary_text += f"资产分类金额Top3: {asset_type_sum_summary}\n"
        # 资产状态金额Top3
        if '资产状态' in df_filtered.columns:
            status_sum = df_filtered.groupby('资产状态')['实际金额'].sum().sort_values(ascending=False).head(3)
            status_sum_summary = ', '.join([f"{idx}: ¥{val:,.2f}" for idx, val in status_sum.items()])
            summary_text += f"资产状态金额Top3: {status_sum_summary}\n"
        # 供应商金额Top3
        supplier_sum = df_filtered.groupby('供应商')['实际金额'].sum().sort_values(ascending=False).head(3)
        supplier_sum_summary = ', '.join([f"{idx}: ¥{val:,.2f}" for idx, val in supplier_sum.items()])
        summary_text += f"供应商金额Top3: {supplier_sum_summary}\n"
        # 设备数量最多的品牌
        if '品牌类别' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录，同时排除品牌类别中包含'费'字的记录
            df_brand_top_filtered = filter_device_count_data(df_filtered)
            df_brand_top_filtered = df_brand_top_filtered[~df_brand_top_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
            if not df_brand_top_filtered.empty:
                top_brand = df_brand_top_filtered['品牌类别'].value_counts().idxmax()
                top_brand_count = df_brand_top_filtered['品牌类别'].value_counts().max()
                summary_text += f"设备数量最多的品牌: {top_brand}（{top_brand_count}台）\n"
        # 设备数量最多的资产分类
        if '资产分类' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录
            df_asset_type_top_filtered = filter_device_count_data(df_filtered)
            top_type = df_asset_type_top_filtered['资产分类'].value_counts().idxmax()
            top_type_count = df_asset_type_top_filtered['资产分类'].value_counts().max()
            summary_text += f"设备数量最多的资产分类: {top_type}（{top_type_count}台）\n"
        # 设备数量最多的资产状态
        if '资产状态' in df_filtered.columns:
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录
            df_status_top_filtered = filter_device_count_data(df_filtered)
            top_status = df_status_top_filtered['资产状态'].value_counts().idxmax()
            top_status_count = df_status_top_filtered['资产状态'].value_counts().max()
            summary_text += f"设备数量最多的资产状态: {top_status}（{top_status_count}台）\n"
        # 设备数量最多的供应商
        top_supplier = df_filtered['供应商'].value_counts().idxmax()
        top_supplier_count = df_filtered['供应商'].value_counts().max()
        summary_text += f"设备数量最多的供应商: {top_supplier}（{top_supplier_count}台）\n"

        import sys
        import subprocess
        try:
            from openai import OpenAI
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openai"])
            from openai import OpenAI
        import os
        api_key = os.environ.get("ARK_API_KEY", "1cfb2a3e-08d5-4aac-a952-b41c14e00a3a")
        ai_result = st.session_state.get('ai_result', "")
        if 'ai_analyzing' not in st.session_state:
            st.session_state['ai_analyzing'] = False
        ai_clicked = st.button("AI整体分析", key="ai_analyze_btn", disabled=st.session_state['ai_analyzing'])
        if ai_clicked:
            st.session_state['ai_analyzing'] = True
            st.session_state['ai_result'] = ""
            st.rerun()
        # 分析中提示
        if st.session_state['ai_analyzing']:
            st.markdown("""
            <div class="data-card">
                <p style="font-size: 16px; margin-bottom: 10px;"><span style="font-weight: bold;">AI智能分析:</span></p>
                <div style="font-size: 15px; color: #2c3e50;">AI分析进行中，请稍候<span class='dot-ani'></span></div>
            </div>
            """, unsafe_allow_html=True)
            # 只在本次分析时执行AI调用
            if not ai_result:
                if not api_key:
                    st.session_state['ai_result'] = "未检测到环境变量 ARK_API_KEY，请配置后重启应用。"
                else:
                    client = OpenAI(
                        base_url="https://ark.cn-beijing.volces.com/api/v3",
                        api_key=api_key,
                    )
                    try:
                        ai_prompt = f"请基于以下IT设备租赁账单多维度统计数据，进行详细分析、风险提示和优化建议，内容尽量丰富，输出问纯文本。请直接输出不用回答好的等。：\n{summary_text}"
                        response = client.chat.completions.create(
                            model="doubao-seed-1-6-250615",
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {"type": "text", "text": ai_prompt}
                                    ],
                                }
                            ],
                        )
                        st.session_state['ai_result'] = response.choices[0].message.content
                    except Exception as e:
                        st.session_state['ai_result'] = f"AI分析异常: {e}"
                st.session_state['ai_analyzing'] = False
                st.rerun()
        else:
            # 分析完成后显示结果或提示
            st.markdown(f"""
            <div class="data-card">
                <p style="font-size: 16px; margin-bottom: 10px;"><span style="font-weight: bold;">AI智能分析:</span></p>
                <div style="font-size: 15px; color: #2c3e50;">{ai_result if ai_result else '请点击“AI整体分析”按钮获取智能分析结果。'}</div>
            </div>
            """, unsafe_allow_html=True)

    # 关键指标概览（对应PDF中的关键指标概览）
    with st.container():
        st.markdown('<h3 class="section-title">📊 关键指标概览</h3>', unsafe_allow_html=True)
        # 四列等宽分布（强迫症对称布局）
        col1, col2, col3, col4 = st.columns(4, gap="large")
        with col1:
            st.markdown("""
                <div class="metric-card">
                    <div class="metric-value">¥{:.2f}</div>
                    <div class="metric-label">总租赁费用</div>
                </div>
            """.format(total_cost), unsafe_allow_html=True)
        with col2:
            st.markdown("""
                <div class="metric-card">
                    <div class="metric-value">{}</div>
                    <div class="metric-label">设备总数</div>
                </div>
            """.format(device_count), unsafe_allow_html=True)
        with col3:
            st.markdown("""
                <div class="metric-card">
                    <div class="metric-value">¥{:.2f}</div>
                    <div class="metric-label">平均设备月租</div>
                </div>
            """.format(avg_monthly), unsafe_allow_html=True)
        with col4:
            st.markdown("""
                <div class="metric-card">
                    <div class="metric-value">{}</div>
                    <div class="metric-label">部门数量</div>
                </div>
            """.format(dept_count), unsafe_allow_html=True)

    # 各平台关键指标（对应PDF中的各平台关键指标）
    with st.container():
        st.markdown('<h3 class="section-title">🏷️ 各平台关键指标</h3>', unsafe_allow_html=True)
        
        # 设备数：排除包含'费'或'赔偿'字段的记录
        df_platform_device_filtered = filter_device_count_data(df_filtered)
        platform_device_data = df_platform_device_filtered.groupby('供应商')['实际金额'].count()
        
        # 总费用：使用所有数据
        platform_cost_data = df_filtered.groupby('供应商')['实际金额'].sum()
        
        # 合并数据
        platform_data = pd.DataFrame({
            '设备数': platform_device_data,
            '总费用': platform_cost_data
        }).fillna(0)
        
        # 平台顺序：易点云、小熊U租、其他
        supplier_order = ['易点云', '小熊U租'] + [s for s in platform_data.index if s not in ['易点云', '小熊U租']]
        platform_data = platform_data.reindex(supplier_order).dropna(how='all')
        
        if not platform_data.empty:
            # 双平台卡片并排（易点云+小熊U租）
            col1, col2 = st.columns(2, gap="large")
            with col1:
                # 易点云数据
                if '易点云' in platform_data.index:
                    yd_device_count = int(platform_data.loc['易点云', '设备数'])
                    yd_total_cost = platform_data.loc['易点云', '总费用']
                    apple_avg = df_filtered[(df_filtered['供应商']=='易点云') & (df_filtered['品牌类别']=='苹果')]['实际金额'].mean()
                    win_avg = df_filtered[(df_filtered['供应商']=='易点云') & (df_filtered['品牌类别']=='Windows')]['实际金额'].mean()
                    st.markdown("""
                        <div class="data-card">
                            <h4 style="font-size: 18px; color: #2c3e50; margin-top: 0;">易点云</h4>
                            <p style="font-size: 16px; margin-bottom: 8px;"><span style="font-weight: bold;">设备数:</span> {}</p>
                            <p style="font-size: 16px; margin-bottom: 8px;"><span style="font-weight: bold;">总费用:</span> ¥{:.2f}</p>
                            <p style="font-size: 16px; margin-bottom: 0;"><span style="font-weight: bold;">平均月租(苹果):</span> ¥{:.2f}</p>
                            <p style="font-size: 16px; margin-top: 2px;"><span style="font-weight: bold;">平均月租(Windows):</span> ¥{:.2f}</p>
                        </div>
                    """.format(yd_device_count, yd_total_cost, apple_avg if not pd.isna(apple_avg) else 0, win_avg if not pd.isna(win_avg) else 0), unsafe_allow_html=True)
                else:
                    st.markdown("<div class='data-card'><h4 style='font-size: 18px; color: #2c3e50; margin-top: 0;'>易点云</h4><p>无数据</p></div>", unsafe_allow_html=True)
            with col2:
                # 小熊U租数据
                if '小熊U租' in platform_data.index:
                    xz_device_count = int(platform_data.loc['小熊U租', '设备数'])
                    xz_total_cost = platform_data.loc['小熊U租', '总费用']
                    apple_avg = df_filtered[(df_filtered['供应商']=='小熊U租') & (df_filtered['品牌类别']=='苹果')]['实际金额'].mean()
                    win_avg = df_filtered[(df_filtered['供应商']=='小熊U租') & (df_filtered['品牌类别']=='Windows')]['实际金额'].mean()
                    st.markdown("""
                        <div class="data-card">
                            <h4 style="font-size: 18px; color: #2c3e50; margin-top: 0;">小熊U租</h4>
                            <p style="font-size: 16px; margin-bottom: 8px;"><span style="font-weight: bold;">设备数:</span> {}</p>
                            <p style="font-size: 16px; margin-bottom: 8px;"><span style="font-weight: bold;">总费用:</span> ¥{:.2f}</p>
                            <p style="font-size: 16px; margin-bottom: 0;"><span style="font-weight: bold;">平均月租(苹果):</span> ¥{:.2f}</p>
                            <p style="font-size: 16px; margin-top: 2px;"><span style="font-weight: bold;">平均月租(Windows):</span> ¥{:.2f}</p>
                        </div>
                    """.format(xz_device_count, xz_total_cost, apple_avg if not pd.isna(apple_avg) else 0, win_avg if not pd.isna(win_avg) else 0), unsafe_allow_html=True)
                else:
                    st.markdown("<div class='data-card'><h4 style='font-size: 18px; color: #2c3e50; margin-top: 0;'>小熊U租</h4><p>无数据</p></div>", unsafe_allow_html=True)
            
            # 添加柱状图和饼图
            if not platform_data.empty:
                # 柱状图
                st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>各平台总费用分布</h4>", unsafe_allow_html=True)
                # 为每个供应商分配不同颜色
                colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75']
                platform_data_reset = platform_data.reset_index()
                platform_data_reset['颜色'] = [colors[i % len(colors)] for i in range(len(platform_data_reset))]
                
                fig_platform_bar = px.bar(
                    platform_data_reset,
                    x='供应商',
                    y='总费用',
                    text='总费用',
                    color='供应商',
                    color_discrete_sequence=colors,
                    title=""
                )
                fig_platform_bar.update_traces(texttemplate='¥%{text:,.0f}', textposition='outside')
                fig_platform_bar.update_layout(
                    yaxis_title='总费用(元)', 
                    xaxis_title='供应商', 
                    margin=dict(t=20, b=10),
                    showlegend=True,
                    legend=dict(
                        orientation="h",
                        yanchor="bottom",
                        y=1.02,
                        xanchor="right",
                        x=1
                    )
                )
                st.plotly_chart(fig_platform_bar, use_container_width=True)
                
                # 饼图
                st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>各平台费用占比</h4>", unsafe_allow_html=True)
                fig_platform_pie = px.pie(
                    platform_data_reset,
                    values='总费用',
                    names='供应商',
                    title="",
                    hole=0.3
                )
                fig_platform_pie.update_traces(textposition='inside', textinfo='percent+label')
                fig_platform_pie.update_layout(margin=dict(t=20, b=10))
                st.plotly_chart(fig_platform_pie, use_container_width=True)
                
                # 各平台关键指标明细表
                st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>各平台关键指标明细表</h4>", unsafe_allow_html=True)
                
                # 准备表格数据
                platform_detail_data = []
                for supplier, data in platform_data.iterrows():
                    # 计算苹果和Windows平均月租
                    apple_avg = df_filtered[(df_filtered['供应商']==supplier) & (df_filtered['品牌类别']=='苹果')]['实际金额'].mean()
                    win_avg = df_filtered[(df_filtered['供应商']==supplier) & (df_filtered['品牌类别']=='Windows')]['实际金额'].mean()
                    
                    platform_detail_data.append({
                        '供应商': supplier,
                        '设备数': int(data["设备数"]),
                        '总费用': f'¥{data["总费用"]:,.2f}',
                        '平均月租(苹果)': f'¥{apple_avg if not pd.isna(apple_avg) else 0:.2f}',
                        '平均月租(Windows)': f'¥{win_avg if not pd.isna(win_avg) else 0:.2f}'
                    })
                
                # 显示表格
                platform_detail_df = pd.DataFrame(platform_detail_data)
                st.dataframe(
                    platform_detail_df, 
                    use_container_width=True, 
                    hide_index=True,
                    column_config={
                        "供应商": st.column_config.TextColumn("供应商", help="供应商名称"),
                        "设备数": st.column_config.NumberColumn("设备数", help="设备总数量", format="%d"),
                        "总费用": st.column_config.TextColumn("总费用", help="总租赁费用"),
                        "平均月租(苹果)": st.column_config.TextColumn("平均月租(苹果)", help="苹果设备平均月租"),
                        "平均月租(Windows)": st.column_config.TextColumn("平均月租(Windows)", help="Windows设备平均月租")
                    }
                )
        # ...区块AI分析已移除...
    # 各平台设备数量（分品牌）分析（独占一行）
    with st.container():
        st.markdown('<h3 class="section-title">🔢 各平台设备数量（分品牌）</h3>', unsafe_allow_html=True)
        # 应用过滤函数排除包含'费'或'赔偿'字段的记录，同时排除品牌类别中包含'费'字的记录
        df_platform_web_filtered = filter_device_count_data(df_filtered)
        df_platform_web_filtered = df_platform_web_filtered[~df_platform_web_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
        platform_brand_devices = df_platform_web_filtered.groupby(['供应商', '品牌类别']).size().reset_index(name='设备数量')
        if not platform_brand_devices.empty:
            # 【颜色示例】可自定义分品牌柱状图配色，如：
            # colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75']
            # 用法：ax.bar(..., color=colors[i % len(colors)])
            # 柱状图标题已通过st.markdown显示在上面
            
            # 柱状图（按固定顺序：易点云、小熊U租、其他）
            suppliers_unique = platform_brand_devices['供应商'].unique()
            supplier_order = ['易点云', '小熊U租'] + [s for s in suppliers_unique if s not in ['易点云', '小熊U租']]
            suppliers = [s for s in supplier_order if s in suppliers_unique]
            brands = platform_brand_devices['品牌类别'].unique()
            
            fig, ax = plt.subplots(figsize=(12, 7))  # 增加图表宽度
            x = np.arange(len(suppliers))
            
            # 根据品牌数量动态调整柱子宽度，避免重叠
            max_width = 0.8  # 所有柱子的最大总宽度
            width = max_width / len(brands) if len(brands) > 0 else 0.35
            width = min(width, 0.35)  # 限制最大宽度
            
            # 为每个品牌分配不同颜色
            colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#F79646', '#2C4D75', '#FF6B6B', '#4ECDC4']
            
            for i, brand in enumerate(brands):
                brand_data = []
                for supplier in suppliers:
                    count = platform_brand_devices[
                        (platform_brand_devices['供应商'] == supplier) & 
                        (platform_brand_devices['品牌类别'] == brand)
                    ]['设备数量'].sum()
                    brand_data.append(count)
                
                # 计算每个品牌柱子的x位置，确保不重叠
                x_pos = x + (i - len(brands)/2 + 0.5) * width
                color = colors[i % len(colors)]
                bars = ax.bar(x_pos, brand_data, width, label=brand, color=color, alpha=0.8)
                
                # 添加数值标签
                for j, bar in enumerate(bars):
                    height = bar.get_height()
                    if height > 0:
                        ax.text(bar.get_x() + bar.get_width()/2., height,
                               f'{int(height)}台', ha='center', va='bottom', fontsize=9)
            
            ax.set_xlabel('供应商', fontsize=12)
            ax.set_ylabel('设备数量', fontsize=12)
            ax.set_xticks(x)
            ax.set_xticklabels(suppliers, rotation=0 if len(suppliers) <= 3 else 15)
            
            # 优化图例位置，避免与柱状图重叠
            if len(brands) <= 4:
                ax.legend(title="品牌类别", loc='upper right', fontsize=10)
            else:
                ax.legend(title="品牌类别", bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=9)
            
            # 设置图表边距，确保图例不被截断
            plt.subplots_adjust(right=0.85 if len(brands) > 4 else 0.95)
            plt.tight_layout()
            # 使用streamlit显示图表
            st.pyplot(fig)
            
            # 饼图（平台总设备占比）
            platform_total = df_platform_web_filtered['供应商'].value_counts().reset_index()
            platform_total.columns = ['供应商', '设备数量']
            platform_total['供应商'] = pd.Categorical(platform_total['供应商'], categories=suppliers, ordered=True)
            platform_total = platform_total.sort_values('供应商')
            
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>各平台设备数量占比</h4>", unsafe_allow_html=True)
            fig_platform_pie = plt.figure(figsize=(8, 6))
            plt.pie(platform_total['设备数量'], labels=platform_total['供应商'], autopct='%1.1f%%', startangle=90, colors=plt.cm.Set3.colors)
            plt.axis('equal')
            # 使用streamlit显示图表
            st.pyplot(fig_platform_pie)
            
            # 明细表格
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>各平台设备数量明细表</h4>", unsafe_allow_html=True)
            st.dataframe(
                platform_brand_devices, 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "供应商": st.column_config.TextColumn("供应商", help="供应商名称"),
                    "品牌类别": st.column_config.TextColumn("品牌类别", help="设备品牌类别"),
                    "设备数量": st.column_config.NumberColumn("设备数量", help="设备总数量", format="%d")
                }
            )
            
            # 添加说明文字
            st.info("📝 说明：品牌分析已排除品牌类别名称中包含'费'字的记录。")
    
    # 供应商平台分析（对应PDF中的供应商平台分析）
    with st.container():
        st.markdown('<h3 class="section-title">🛒 供应商平台分析</h3>', unsafe_allow_html=True)
        
        # 排除品牌类别中包含'费'字的记录
        df_supplier_web_filtered = df_filtered[~df_filtered['品牌类别'].astype(str).str.contains('费', na=False)]
        platform_cost = df_supplier_web_filtered.groupby('供应商')['实际金额'].sum()
        # 平台顺序：易点云、小熊U租、其他
        supplier_order = ['易点云', '小熊U租'] + [s for s in platform_cost.index if s not in ['易点云', '小熊U租']]
        platform_cost = platform_cost.reindex(supplier_order).dropna(how='all')
        # 表格在柱状图下，简化为每品牌总价格、数量（共6列）
        
        brand_detail = df_supplier_web_filtered.groupby(['供应商', '品牌类别'])['实际金额'].agg(['sum', 'count', 'mean']).reset_index()
        brand_detail.columns = ['供应商', '品牌类别', '总价', '数量', '均价']
        brand_detail['总价'] = brand_detail['总价'].apply(lambda x: f'¥{x:,.2f}')
        brand_detail['均价'] = brand_detail['均价'].apply(lambda x: f'¥{x:,.2f}')
        # 明细表格排序
        brand_detail['供应商'] = pd.Categorical(brand_detail['供应商'], categories=supplier_order, ordered=True)
        brand_detail = brand_detail.sort_values('供应商')
        # 平台费用明细（分类别）柱状图和饼图上下排列
        brand_detail_bar = df_supplier_web_filtered.groupby(['供应商', '品牌类别'])['实际金额'].sum().reset_index()
        # 确保供应商顺序：易点云、小熊U租、其他
        brand_detail_bar['供应商'] = pd.Categorical(brand_detail_bar['供应商'], categories=supplier_order, ordered=True)
        brand_detail_bar = brand_detail_bar.sort_values('供应商')
        st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>平台-类别费用分布</h4>", unsafe_allow_html=True)
        fig_brand_bar = px.bar(
            brand_detail_bar,
            x='供应商',
            y='实际金额',
            color='品牌类别',
            barmode='group',
            text='实际金额',
            title=""
        )
        fig_brand_bar.update_traces(texttemplate='¥%{text:,.0f}', textposition='outside')
        fig_brand_bar.update_layout(yaxis_title='费用(元)', xaxis_title='供应商', margin=dict(t=20, b=10))
        st.plotly_chart(fig_brand_bar, use_container_width=True)
        st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>平台-类别费用占比</h4>", unsafe_allow_html=True)
        fig_brand_pie = px.pie(
            brand_detail_bar,
            values='实际金额',
            names='品牌类别',
            title="",
            hole=0.3,
            color='品牌类别'
        )
        fig_brand_pie.update_traces(textposition='inside', textinfo='percent+label')
        fig_brand_pie.update_layout(margin=dict(t=20, b=10))
        st.plotly_chart(fig_brand_pie, use_container_width=True)
        # 明细表格放到饼图下方
        st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>平台费用明细（分类别）</h4>", unsafe_allow_html=True)
        st.dataframe(
            brand_detail, 
            use_container_width=True, 
            hide_index=True,
            column_config={
                "供应商": st.column_config.TextColumn("供应商", help="供应商名称"),
                "品牌类别": st.column_config.TextColumn("类别", help="设备类别"),
                "总价": st.column_config.TextColumn("总价", help="总租赁费用"),
                "数量": st.column_config.NumberColumn("数量", help="设备数量", format="%d"),
                "均价": st.column_config.TextColumn("均价", help="平均单价")
            }
        )
        
        # 添加说明文字
        st.info("📝 说明：供应商平台分析已排除品牌类别名称中包含'费'字的记录。")
    
    # 资产状态分析
    if '资产状态' in df_filtered.columns:
        with st.container():
            st.markdown('<h3 class="section-title">📦 资产状态分析</h3>', unsafe_allow_html=True)
            
            # 状态分类映射函数
            def categorize_status(status):
                status_str = str(status).strip()
                if '领用' in status_str:
                    return '领用'
                elif '空闲' in status_str:
                    return '空闲'
                elif '退租' in status_str or '处置' in status_str:
                    return '已处置'
                elif '费' in status_str or '运费' in status_str or '赔偿' in status_str:
                    return '费用'
                else:
                    # 其他状态根据内容判断
                    if any(keyword in status_str for keyword in ['归还', '返还', '报废', '损坏']):
                        return '已处置'
                    else:
                        return '领用'  # 默认归类为领用
            
            # 计算费用总额（使用所有数据）
            status_sum = df_filtered.groupby('资产状态')['实际金额'].sum().reset_index()
            status_sum.columns = ['资产状态', '总金额']
            status_sum['状态分类'] = status_sum['资产状态'].apply(categorize_status)
            
            # 计算设备数量（排除费用类状态）
            df_status_device_filtered = df_filtered[~df_filtered['资产状态'].apply(lambda x: categorize_status(x) == '费用')]
            status_count = df_status_device_filtered.groupby('资产状态').size().reset_index()
            status_count.columns = ['资产状态', '设备数量']
            status_count['状态分类'] = status_count['资产状态'].apply(categorize_status)
            
            # 按状态分类汇总
            category_sum = status_sum.groupby('状态分类')['总金额'].sum().reset_index()
            category_count = status_count.groupby('状态分类')['设备数量'].sum().reset_index()
            
            # 合并数据
            category_data = pd.merge(category_sum, category_count, on='状态分类', how='outer')
            category_data['设备数量'] = category_data['设备数量'].fillna(0).astype(int)
            category_data['总金额'] = category_data['总金额'].fillna(0)
            
            # 按指定顺序排序：领用-空闲-已处置-费用
            order_map = {'领用': 1, '空闲': 2, '已处置': 3, '费用': 4}
            category_data['排序'] = category_data['状态分类'].map(order_map)
            category_data = category_data.sort_values('排序').reset_index(drop=True)
            
            # 柱状图 - 显示各状态分类设备数量分布
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 0;'>资产状态分类设备数量分布</h4>", unsafe_allow_html=True)
            
            # 创建柱状图
            fig_status_bar = px.bar(
                category_data[category_data['状态分类'] != '费用'],  # 费用类不显示设备数量
                x='状态分类',
                y='设备数量',
                text='设备数量',
                title="",
                color='状态分类',
                color_discrete_map={
                    '领用': '#2E75B6',
                    '空闲': '#9BBB59', 
                    '已处置': '#C0504D',
                    '费用': '#E46C0A'
                }
            )
            fig_status_bar.update_traces(texttemplate='%{text}', textposition='outside')
            fig_status_bar.update_layout(
                yaxis_title='设备数量', 
                xaxis_title='资产状态分类', 
                margin=dict(t=20, b=10), 
                showlegend=True,
                legend=dict(
                    title="状态分类",
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1
                )
            )
            st.plotly_chart(fig_status_bar, use_container_width=True)
            
            # 饼图 - 按状态分类显示
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>资产状态分类占比</h4>", unsafe_allow_html=True)
            
            # 只显示有设备数量的分类（排除费用类）
            pie_data = category_data[
                (category_data['设备数量'] > 0) & 
                (category_data['状态分类'] != '费用')
            ]
            
            if not pie_data.empty:
                fig_type_pie = px.pie(
                    pie_data,
                    values='设备数量',
                    names='状态分类',
                    title="",
                    hole=0.3,
                    color='状态分类',
                    color_discrete_map={
                        '领用': '#2E75B6',
                        '空闲': '#9BBB59', 
                        '已处置': '#C0504D',
                        '费用': '#E46C0A'
                    }
                )
                fig_type_pie.update_traces(textposition='inside', textinfo='percent+label')
                fig_type_pie.update_layout(margin=dict(t=20, b=10))
                st.plotly_chart(fig_type_pie, use_container_width=True)
            
            # 汇总统计表
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>资产状态分类汇总</h4>", unsafe_allow_html=True)
            
            # 创建汇总表格
            summary_data = []
            for _, row in category_data.iterrows():
                summary_data.append({
                    '状态分类': row['状态分类'],
                    '设备数量': row['设备数量'] if row['状态分类'] != '费用' else '-',
                    '总金额': f'¥{row["总金额"]:,.2f}',
                    '说明': '不统计设备数量' if row['状态分类'] == '费用' else f'{row["设备数量"]}台设备'
                })
            
            summary_df = pd.DataFrame(summary_data)
            st.dataframe(
                summary_df, 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "状态分类": st.column_config.TextColumn("状态分类", help="状态分类"),
                    "设备数量": st.column_config.TextColumn("设备数量", help="设备数量统计"),
                    "总金额": st.column_config.TextColumn("总金额", help="总费用"),
                    "说明": st.column_config.TextColumn("说明", help="备注信息")
                }
            )
            
            # 添加说明文字
            st.info("📝 说明：资产状态按领用、空闲、已处置、费用四类展示。设备数量统计不包含费用类状态，费用统计包含所有记录。")
    
    # 资产分类分析
    if '资产分类' in df_filtered.columns:
        with st.container():
            st.markdown('<h3 class="section-title">🏷️ 资产分类分析</h3>', unsafe_allow_html=True)
            # 排除资产分类中包含'费'字的记录
            df_asset_type_web_filtered = df_filtered[~df_filtered['资产分类'].astype(str).str.contains('费', na=False)]
            asset_type_count = df_asset_type_web_filtered['资产分类'].astype(str).value_counts().reset_index()
            asset_type_count.columns = ['资产分类', '设备数量']
            # 对于总金额统计，也排除带有'费'字的资产分类
            asset_type_sum = df_asset_type_web_filtered.groupby('资产分类')['实际金额'].sum().reset_index()
            asset_type_sum.columns = ['资产分类', '总金额']
            asset_type_sum['总金额'] = asset_type_sum['总金额'].apply(lambda x: f'¥{x:,.2f}')
            # 合并数量和金额
            asset_type_table = pd.merge(asset_type_count, asset_type_sum, on='资产分类', how='left')
            # 柱状图
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 0;'>各资产分类设备数量</h4>", unsafe_allow_html=True)
            fig_asset_type_bar = px.bar(
                asset_type_count,
                x='资产分类',
                y='设备数量',
                text='设备数量',
                color='资产分类',
                title="",
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            fig_asset_type_bar.update_traces(texttemplate='%{text}台', textposition='outside')
            fig_asset_type_bar.update_layout(
                yaxis_title='设备数量', 
                xaxis_title='资产分类', 
                margin=dict(t=20, b=10), 
                showlegend=True,
                legend=dict(
                    title="资产分类",
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1
                )
            )
            st.plotly_chart(fig_asset_type_bar, use_container_width=True)
            # 饼图
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>资产分类占比</h4>", unsafe_allow_html=True)
            fig_asset_type_pie = px.pie(
                asset_type_count,
                values='设备数量',
                names='资产分类',
                title="",
                hole=0.3,
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            fig_asset_type_pie.update_traces(textposition='inside', textinfo='percent+label')
            fig_asset_type_pie.update_layout(
                margin=dict(t=20, b=10),
                showlegend=True,
                legend=dict(
                    title="资产分类",
                    orientation="v",
                    yanchor="middle",
                    y=0.5,
                    xanchor="left",
                    x=1.02
                )
            )
            st.plotly_chart(fig_asset_type_pie, use_container_width=True)
            # 明细表格
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>资产分类明细</h4>", unsafe_allow_html=True)
            st.dataframe(
                asset_type_table, 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "资产分类": st.column_config.TextColumn("资产分类", help="设备资产分类"),
                    "设备数量": st.column_config.NumberColumn("设备数量", help="设备数量统计", format="%d"),
                    "总金额": st.column_config.TextColumn("总金额", help="总租赁费用")
                }
            )
            
            # 添加说明文字
            st.info("📝 说明：资产分类分析已排除资产分类名称中包含'费'字的记录。")

    # 7. 部门费用分析（图表在上，表格在下，顺序与web一致）
    with st.container():
        st.markdown('<h3 class="section-title">👥 部门费用分析</h3>', unsafe_allow_html=True)
        dept_cost = df_filtered.groupby('一级部门')['实际金额'].sum().sort_values(ascending=False)
        # 柱状图独占一行
        st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 0;'>部门费用分布</h4>", unsafe_allow_html=True)
        fig_bar = px.bar(
            dept_cost.reset_index(),
            x='一级部门',
            y='实际金额',
            title="",
            text='实际金额',
            color='一级部门',
        )
        fig_bar.update_traces(texttemplate='¥%{text:,.0f}', textposition='outside')
        fig_bar.update_layout(yaxis_title='费用(元)', xaxis_title='部门', margin=dict(t=20, b=10), showlegend=False)
        st.plotly_chart(fig_bar, use_container_width=True)
        # 饼状图
        st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>部门费用占比</h4>", unsafe_allow_html=True)
        fig_pie = px.pie(
            dept_cost.reset_index(),
            values='实际金额',
            names='一级部门',
            title="",
            hole=0.3
        )
        fig_pie.update_traces(textposition='inside', textinfo='percent+label')
        fig_pie.update_layout(margin=dict(t=20, b=10))
        st.plotly_chart(fig_pie, use_container_width=True)
        # 表格在图下方
        st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>部门费用详情</h4>", unsafe_allow_html=True)
        dept_table = dept_cost.reset_index()
        dept_table.columns = ['部门', '总费用']
        dept_table['总费用'] = dept_table['总费用'].apply(lambda x: f'¥{x:,.2f}')
        st.dataframe(
            dept_table, 
            use_container_width=True, 
            hide_index=True,
            column_config={
                "部门": st.column_config.TextColumn("部门", help="一级部门名称"),
                "总费用": st.column_config.TextColumn("总费用", help="部门总租赁费用")
            }
        )
    
    # 8. 人员费用分析（展示费用超500元人员，及其设备数量、类型、金额、部门）
    if '领用人' in df_filtered.columns and '人员编号' in df_filtered.columns:
        with st.container():
            st.markdown('<h3 class="section-title">👤 人员费用分析</h3>', unsafe_allow_html=True)
            # 应用过滤函数排除包含'费'或'赔偿'字段的记录
            df_person_web_filtered = filter_device_count_data(df_filtered)
            # 以（领用人, 人员编号）为唯一标识统计
            person_group = df_person_web_filtered.groupby(['领用人', '人员编号', '一级部门']).agg({
                '实际金额': 'sum',
                '品牌类别': lambda x: ', '.join(sorted(set(map(str, x)))) if len(x) else '',
                '领用人': 'count'
            }).rename(columns={'实际金额': '总费用', '品牌类别': '设备类型', '领用人': '设备数量'}).reset_index()
            # 只保留总费用大于500元的人员
            person_group = person_group[person_group['总费用'] > 500]
            # 按总费用降序排列
            person_group = person_group.sort_values('总费用', ascending=False)
            # 合成唯一标识列
            person_group['人员'] = person_group['领用人'].astype(str) + '（' + person_group['人员编号'].astype(str) + '）'
            # 柱状图
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 0;'>人员费用分布（总费用>500元）</h4>", unsafe_allow_html=True)
            fig_person_bar = px.bar(
                person_group,
                x='人员',
                y='总费用',
                text='总费用',
                color='人员',
                title=""
            )
            fig_person_bar.update_traces(texttemplate='¥%{text:,.0f}', textposition='outside')
            fig_person_bar.update_layout(yaxis_title='费用(元)', xaxis_title='人员', margin=dict(t=20, b=10), showlegend=False)
            st.plotly_chart(fig_person_bar, use_container_width=True)
            # 饼图
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>人员费用占比（总费用>500元）</h4>", unsafe_allow_html=True)
            fig_person_pie = px.pie(
                person_group,
                values='总费用',
                names='人员',
                title="",
                hole=0.3
            )
            fig_person_pie.update_traces(textposition='inside', textinfo='percent+label')
            fig_person_pie.update_layout(margin=dict(t=20, b=10))
            st.plotly_chart(fig_person_pie, use_container_width=True)
            # 详细表格，增加部门字段
            st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 20px;'>人员费用明细（总费用>500元）</h4>", unsafe_allow_html=True)
            person_group['总费用'] = person_group['总费用'].apply(lambda x: f'¥{x:,.2f}')
            st.dataframe(
                person_group[['人员', '一级部门', '设备数量', '设备类型', '总费用']], 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "人员": st.column_config.TextColumn("人员", help="领用人姓名"),
                    "一级部门": st.column_config.TextColumn("一级部门", help="所属一级部门"),
                    "设备数量": st.column_config.NumberColumn("设备数量", help="领用设备数量", format="%d"),
                    "设备类型": st.column_config.TextColumn("设备类型", help="设备品牌类型"),
                    "总费用": st.column_config.TextColumn("总费用", help="总租赁费用")
                }
            )

            # 新增：人员领取多台设备分析（受筛选条件影响）
            if '领用人' in df_filtered.columns and '人员编号' in df_filtered.columns:
                # 应用过滤函数排除包含'费'或'赔偿'字段的记录
                df_multi_web_filtered = filter_device_count_data(df_filtered)
                multi_device_all = df_multi_web_filtered.groupby(['领用人', '人员编号', '一级部门']).agg({
                    '实际金额': 'sum',
                    '品牌类别': lambda x: ', '.join(sorted(set(map(str, x)))) if len(x) else '',
                    '领用人': 'count'
                }).rename(columns={'实际金额': '总费用', '品牌类别': '设备类型', '领用人': '设备数量'}).reset_index()
                multi_device_all = multi_device_all[multi_device_all['设备数量'] > 1]
                multi_device_all = multi_device_all.sort_values('设备数量', ascending=False)
                multi_device_all['人员'] = multi_device_all['领用人'].astype(str) + '（' + multi_device_all['人员编号'].astype(str) + '）'
                if not multi_device_all.empty:
                    st.markdown("<h4 style='font-size: 18px; color: #2c3e50; margin-top: 30px;'>人员领取多台设备分析</h4>", unsafe_allow_html=True)
                    # 柱状图
                    fig_multi_bar = px.bar(
                        multi_device_all,
                        x='人员',
                        y='设备数量',
                        color='一级部门',
                        text='设备数量',
                        title=""
                    )
                    fig_multi_bar.update_traces(texttemplate='%{text}台', textposition='outside')
                    fig_multi_bar.update_layout(yaxis_title='设备数量', xaxis_title='人员', margin=dict(t=20, b=10), showlegend=True)
                    st.plotly_chart(fig_multi_bar, use_container_width=True)
                    # 明细表格
                    multi_device_all['总费用_排序'] = multi_device_all['总费用'].replace({'¥': '', ',': ''}, regex=True).astype(float)
                    multi_device_all = multi_device_all.sort_values('总费用_排序', ascending=False)
                    st.dataframe(
                        multi_device_all[['人员', '一级部门', '设备数量', '设备类型', '总费用']], 
                        use_container_width=True, 
                        hide_index=True,
                        column_config={
                            "人员": st.column_config.TextColumn("人员", help="领用人姓名"),
                            "一级部门": st.column_config.TextColumn("一级部门", help="所属一级部门"),
                            "设备数量": st.column_config.NumberColumn("设备数量", help="领用设备数量", format="%d"),
                            "设备类型": st.column_config.TextColumn("设备类型", help="设备品牌类型"),
                            "总费用": st.column_config.TextColumn("总费用", help="总租赁费用")
                        }
                    )
    

    
if __name__ == "__main__":
    main()
